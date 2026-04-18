"""
AKADEMIK IŞLER BOT — ŞAHSY WERSIÝA (tölegsiz)
Bot → TÜRKMEN dilinde | Faýl → HEMIŞE RUS dilinde
Render.com | Python 3.11

🔧 DÜZEDILEN WERSIÝA:
- Accept-Encoding: identity → "invalid distance too far back" çözüldi
- Şablon global cache → çalt we durnukly
- F.video handler dubl düzedildi
- Token/API key diňe env-den
- Redis URL env-den
"""

import asyncio, base64, copy, io, json, logging, os, re
import httpx
from aiogram import Bot, Dispatcher, F, Router
from aiogram.filters import CommandStart
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
try:
    from aiogram.fsm.storage.redis import RedisStorage
    _HAS_REDIS = True
except ImportError:
    _HAS_REDIS = False
from aiogram.types import (
    BufferedInputFile, CallbackQuery,
    InlineKeyboardButton, InlineKeyboardMarkup, Message,
)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

logging.basicConfig(level=logging.INFO, format="%(asctime)s | %(levelname)s | %(message)s")
log = logging.getLogger(__name__)

# ── SAZLAMALAR ──────────────────────────────────────────────
# Render'de "Environment Variables" bölümünde HÖKMAN goý:
#   BOT_TOKEN        = siziň bot tokeniniz
#   DEEPSEEK_API_KEY = siziň DeepSeek açaryňyz
#   REDIS_URL        = (islege bagly) redis://...
BOT_TOKEN        = os.getenv("BOT_TOKEN")
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
REDIS_URL        = os.getenv("REDIS_URL")

if not BOT_TOKEN or not DEEPSEEK_API_KEY:
    raise RuntimeError(
        "❌ BOT_TOKEN we DEEPSEEK_API_KEY env-de goýulmaly!\n"
        "Render dashboard → Environment Variables bölümüne git we goý."
    )

ADMIN_IDS        = [8512644114, 7404431806]
INTRO_VIDEO_URL  = "https://youtu.be/FX7MlvKpGqA?si=gsmJpuFiQ_gHKFN8"
DEEPSEEK_URL     = "https://api.deepseek.com/v1/chat/completions"
DEEPSEEK_MODEL   = "deepseek-chat"
CARD_NUMBER      = "2202 2084 5873 0067"
PHONE_NUMBER     = "+7 922 309 80 64"
CARD_HOLDER      = "Мекан Н"
PRICE            = {"referat": 300, "doklad": 300, "zadaniye": 150}


TEMPLATE_B64 = (
    "UEsDBBQABgAIAAAAIQBKvAJxbQEAACgGAAATAAgCW0NvbnRlbnRfVHlwZXNdLnhtbCCiBAIooAACAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAC0lMtqwzAQRfeF/oPRtthKuiilxMmij2UbaPoBijRO"
    "RGVJSJPX33ccO6aUJC5NvTHYM/feM2OY0WRbmmQNIWpnczbMBiwBK53SdpGzj9lLes+SiMIqYZyFnO0g"
    "ssn4+mo023mICaltzNkS0T9wHuUSShEz58FSpXChFEivYcG9kJ9iAfx2MLjj0lkEiylWHmw8eoJCrAwm"
    "z1v6XJMEMJElj3VjlZUz4b3RUiDV+dqqHylpk5CRct8Tl9rHG2pg/GhCVTkd0OjeaDVBK0imIuCrKKmL"
    "b1xQXDm5KkmZnbc5wumKQkto9ZWbD05CjLTz0mRtpRTaHvhPckTcGYj/T1H7dscDIgn6AGicOxE2MH/v"
    "jeKbeSdI4Rxah338jda6EwKs6onh4PyLPVCimBvoYw+NdScE0iWC+jm8mGNvcy6SOqfB+UiXLfxh7MPp"
    "qtQpDewhoD6/6TaRrC+eD6qrqEAdyeb7Oz/+AgAA//8DAFBLAwQUAAYACAAAACEAHpEat+8AAABOAgAA"
    "CwAIAl9yZWxzLy5yZWxzIKIEAiigAAIAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAKySwWrD"
    "MAxA74P9g9G9UdrBGKNOL2PQ2xjZBwhbSUwT29hq1/79PNjYAl3pYUfL0tOT0HpznEZ14JRd8BqWVQ2K"
    "vQnW+V7DW/u8eACVhbylMXjWcOIMm+b2Zv3KI0kpyoOLWRWKzxoGkfiImM3AE+UqRPblpwtpIinP1GMk"
    "s6OecVXX95h+M6CZMdXWakhbeweqPUW+hh26zhl+CmY/sZczLZCPwt6yXcRU6pO4Mo1qKfUsGmwwLyWc"
    "kWKsChrwvNHqeqO/p8WJhSwJoQmJL/t8ZlwSWv7niuYZPzbvIVm0X+FvG5xdQfMBAAD//wMAUEsDBBQA"
    "BgAIAAAAIQCVqf+Z1gsAAIWHAAARAAAAd29yZC9kb2N1bWVudC54bWzsXVtv28gVfi/Q/yDoqX1wxPtF"
    "WGdB8bIbIC2CJNs+LmiKsthIokpSlt2nxN5b4bQG2reii2TRRV8WBRxnvfEmTgLsLyD/Uc8ZkrJkyTZF"
    "y7KsUA+8zOXM5XznzJkzM9RHH2+2W6UN2/Mdt7Napm9R5ZLdsdy601lfLX/20FiRyiU/MDt1s+V27NXy"
    "lu2XP77961991K/WXavXtjtBCUh0/Gq/a62Wm0HQrVYqvtW026Z/q+1Ynuu7jeCW5bYrbqPhWHal73r1"
    "CkPRFHnqeq5l+z6Up5qdDdMvJ+SszWzU6p7Zh8xIkKtYTdML7M0TGvTURPiKXJHGCTE5CEELGXqcFDs1"
    "KaGCtRojxOUiBLUao8TnozShcUI+Ssw4JTEfJXackpSP0hic2uMAd7t2ByIbrtc2A3j11itt03vU664A"
    "4a4ZOGtOywm2gCYlpGRMp/MoR40g14BCm61PTUGstN263WLrKRV3tdzzOtUk/8ogP1a9GudPboMcditb"
    "sVCcXLE3g5YfpHm9LH0XZ9cSxUJ6reLZLehHt+M3ne5AO7TzUoPIZkpk47wO2Gi30nT9Lp1R1M5SbVrM"
    "hhOCWaqf8K7dimt+PkWaysBNJDHIkaUKo2WmNWkDgk8KztU1Q51LZ1Q+KQFmjIBg2RkHi5SGlNCoWCfS"
    "jXScjGKV0om5gnSck46lM+rA05UZIuDXg3pzKipM2q8VzGsGZtP0B0BHivZ0leIH5LbaQ33UXb+cIHzi"
    "ub3uCTXnctTunKjEPhonU9BKBGpYyP3LVeZB0+yCpmxb1TvrHdcz11pQIxCPEiC8RDiAVwAK3sijvUnC"
    "kdcl1DHl22BVrbn1Lbx3IY6rdk3PvAOgFFhVYxWDKZNQGJMCDBWTH4RWwYKr318tU5ShsVKNGgTd8zBQ"
    "kQSBkwaBmt0we61gPPm9oSBSi3se3vyuaUETIZHZCGykiBlaDnY6ww1e7vewzWYvcMsVzPYnCyI2TBg0"
    "LFDAtheHejFNz3A7gY8kfcsBIDx02rZf+r3dL91322YHadqmHyi+Y06MbCodf3I2yx8PJiWvkWvLJC0h"
    "9fJ6K/c/w8hKUq/KoM1rrvsIx/IHARgBkB51A2lpx2xDKz/vN5qNrXVOaAiCRUp1W4bj+dipyetdM36L"
    "y07I6Z36gFjcHRM5tfAd1a8Gt8N/h0fh2/AoehJth4fRY3I/CN+Vwnfhi+hxuB++gqcDuGOivVJ4VIKn"
    "/WgnfI2Pz8N3kOMJ0PgZrq8h6c+l8PvwMPwRiUG6ryDqCEsM4nJj9oxLB8/yNU6VhEI6LsP0jHKxiHi9"
    "CKmo4wmbgAFdz/Ztb8Mu3x7GWvgmegrYfBcelsKXBJc7ELU/wPQhRCbR++EBBL0jr8dp4CnAozwMkYx2"
    "oq8h+jD8CctDaYDADMAWKZmVDdoogF0A+wwVfBDtAkS/AVS9PFPxZkCaoHAqrwtqgbS5GRiLiLdsA/9k"
    "dfrLD+EzokyPyWh+hKM5jvY4iqO2THRstIsRaDM8Cd8QBTyiMQd5M4CW0lWDo7hCPRagzaIq3wPS3gDy"
    "0Fb9kmjGr+EpxRsM+DguH8SGbJJs+5c3GXDIGjqjGRpX4PAybPT/ktaHkdIQFaswEpYJp2dYU1qNVlSl"
    "YNOisGkx1ckMGggaw0StNKo0grVWckvSrLX+CCT6YH1xDI91Cra6wPv6pnke6yHbXXPL7aFTIs7QcDbt"
    "+kmk6z5K81HwQ8IN9Ezcd/sp6MzhNxKpuq1euzMUPxLQcT+tmZ2BB8T9Q/oWN3K4SZ94Th0f1+EONOIG"
    "MpQkxhUcCeZYlnRnTCLNGQxQMSJbQ5JMi5LKy9oF7rCHBFMSIzAsQzAVpMzxPrWd9SZ2YMxTXji7w7Fu"
    "aT4rvqZvCfdI405zD7MlCUe1EMOxMqXJOSzt8xXOUDsT4QIM+sk9bVbLbgSYq+sC+mWZTWuapMypoq5L"
    "Vsf1Sddz3YbuIYBibvhdu9UiLrykqotQX2KNfA8G8WswOtDmJabGiKFxVlP0DhH0WKUQjJ0DTBAvOTsw"
    "2Zpa02m2VgDzwwbmPwGUx4nrYBtdYPGcLTc8iQIltDHx+SbWJGXPyAIj1WTtfFxmV/YiM3NlD3SUlrPe"
    "mUxworSJOq9LkoC0rljaCtmZn+z8C+TlC3RqoBNuTgo9B/iomsLqbCpoVzwTuvGIXHIP8v/I0lu8mnEE"
    "z/u4TgdjwHviSD4gTrp4xYIEQbJoN/omCcah4TiNOApf4fIfoD9ZCAyPR2RgpsMCL3OCYrCIwlkMC6LE"
    "z28OwHEyo3H0Ms0BHLKu7WGXrpZXaCqZJhejSI5R5FuQuzHpm2Y0WaDGVCdogDFxMHSKVnilEIdCHCaI"
    "wzMcXMCsOsK58hKJQS5r78wxRWB5Q2X0HEI0E5tu6ef0y20FstItir1F0aVffhg3CDMtggk1lRdqDFvg"
    "r8Df9PgD2P03PIq+IPMI3A6wV4r+lnVikgGdlKGotKHgKkuBzgKdU6ITd6/Ek1zcr4L3o2gPprh4/bKE"
    "3lKYEONWWNw4uE0wTKAZ7U1UncnAj7e4gEGnTDUDpnhV5xjxAo17ratgeXxTYk1VVVm/JkldDnsfKpRd"
    "KC9hjeZxfNMUI9LKfI4yfCi+x0l6BW/xjoNTu+U5TdQ1ptgtfym2XP1uJaVGK5pUbG5cbDZRGq+yFJt7"
    "2i1JDK2jMpw9m1jhguGKbPy5S9JKPH2agWtu0JyWfbl3b66p/qz5NfkIF42VSY5wrW+IzJ85y0JbdVLF"
    "aU9wJV3mTeTozes4Yuk+Dw9xBZOcv9omG7hLZGv2cbRTHXMOnNpAG/2VZHxRAjsY8r6CWdrL6HG0B09v"
    "U4MYjyi8hsCd8Cco4e+YBfM+iW3raDeTr4HhaoymijlsxGnl66KJGy/Kg01ws/AYI3GwtXg2scWXeYBT"
    "VVpiLtrEVzBxoZnICCqnsGwOv0rBxIVhokDVeEPUZ74XqWDiHJnIGgbH8sbMt28WTJwjE2lBFRVJziGJ"
    "M5n0FZy9unkiJ9cERbyu6fx1c/ZkWkl8pEvE7v4SH2KafNC3NOUv/Ee0mxy+xOW6N1m2BPEyT3MMvdyj"
    "GbIWOODgtxZlecwBU0jKTZeU8Ltom3xI5DB8ix4R9IbsgCi8j3ZHz7enlxveC7iZ+D/hs/DbFYZdYcIX"
    "GQSd1gTakI3c20ZYmZJikSuk//oFfLF3Fc4O5t9FO2QTylSnopak9WcYBZl6YvmQ8Bzsmlfh8cmHTKZB"
    "xBm2D1PTFVGez16lwiAqDKIrkYzPh38ZDAFRZAydvbaPqMwP9QuG53ij0BDTk+XOEaYnYTcY1TNrJgL5"
    "N2Q6+yNc8TtWT3H5csjOD/d/mwHvFMtxHE3lwHvh1bsGXT2ZiYIq6oasYYcUTLypTOR1jpH4WrH8dZOZ"
    "yMo0VROu7eNeBWevTjxptiZReb6UUaycFNOfxfIHk6Os8TcmyUpIpo9AMzVZVaULvn9TTIwKpC/ARD/L"
    "TJ9XdZnljKXaM1ZM6pduUp+c94sD8bxp+o3qvSzTe1FWaNGgLvg8TQHyRQL5ZE5yLGgrJs/4W3BysThJ"
    "sYJGsSr2SsHJm81JieVlg1kq5+kHykmVETSZyrH/o+Dkgo2TMm/oqpLD4inmqTdBUClKqPF0jmPjcz96"
    "nOfk6qy5snaaK1fjHKVVXlN1cQ77ZAuuZOcKz6sizdNz8HAUXJmCK7REKQoZYq5jgCpYNYlV/dOOJVUU"
    "dEUgHbcEzUQ/UfgMzz7gf0AmfyOJXzcL30VfkdfX4f6oB/WsH7FcYspnIpyiaVHl52BML+afmF25CuFk"
    "rcYzau6tmRLFKxx6kj60Xr9Q7Bfon+RmJ/bxvyE+LTEUMyK96WU5GipM1Ey+bQX3BkzOKBEPIBNBg8Iz"
    "egL99QdYzT7YurRMxXuE8aCsFP91FCb4nYnlBG4X07DoFEq/xCbxpDwYOQO3fRIbr7HTIkW+BtO0zTqO"
    "xyLx1lcbrkuG5+R1vRckozUpzXJb2EPJ0i6mIcF118K/r0LaMIbfcwKrGX+AJ+6ruDPIY/wH7xWSpdcG"
    "Mb79fwAAAP//AwBQSwMEFAAGAAgAAAAhABEXoNkNAQAAOQQAABwACAF3b3JkL19yZWxzL2RvY3VtZW50"
    "LnhtbC5yZWxzIKIEASigAAEAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAArJPNTsMwEITvSLyDtXfipEBBqE4vqFKvEB7A"
    "TTY/IllH9hbI22MVpaRQRT34uGPtzOexvFp/da34QOsaQwqSKAaBlJuioUrBW7a5eQThWFOhW0OoYEAH"
    "6/T6avWCrWa/5Oqmd8K7kFNQM/dPUrq8xk67yPRI/qQ0ttPsR1vJXufvukK5iOOltFMPSE88xbZQYLfF"
    "LYhs6PESb1OWTY7PJt93SHwmQn7i7hWZ/eWct9W2QlYwESPvCPI8yENIEPa7+ItwGH/EZI5hEZLB/Wti"
    "VOYQkqAIPLQ4BTjMc/HLkPGlIc70rp28xFGag7gPCYFUkOFpC6Myh3AXtgfDfxiO0gghTz58+g0AAP//"
    "AwBQSwMEFAAGAAgAAAAhAK7IR2fCAgAAzQsAABIAAAB3b3JkL2Zvb3Rub3Rlcy54bWzUlktvozAQgO8r"
    "7X9A3FMDSZOAmlTqZrvqrWp3f4BrTEDFD9kmJP9+xzzT0q0IPS0HHmPP55nxzOCb2yPLnQNVOhN84/pX"
    "nutQTkSc8f3G/fP7frZ2HW0wj3EuON24J6rd2+33bzdllAhhuDBUO8DgOool2bipMTJCSJOUMqyvWEaU"
    "0CIxV0QwJJIkIxSVQsUo8HyvepNKEKo1LPgD8wPWboMjx3G0WOESlC1wgUiKlaHHnuFfDLlGIVoPQcEE"
    "EHgY+EPU/GLUElmrBqDFJBBYNSBdTyN94NxyGikYklbTSPMhaT2NNEgnNkxwISmHwUQohg18qj1iWL0W"
    "cgZgiU32kuWZOQHTW7YYnPHXCRaBVkdg8/hiwgoxEdN8HrcUsXELxaNGf9bpW9OjWr95dBo0H7csLBci"
    "ejS5Nq2uGhO7Wn0nSMEoN1XUkKI5xFFwnWay6w5sKg0G0xZy+CwAB5a380rpjyy1f7W2Xb0NPXCM+c3e"
    "sby2/HOi743YTYvoNMaY8HbN1hIGGdwvPCk0Z8H1RzafFhAMAEtCR/4sWsa6YSDSV7flZCPLquXUu2I5"
    "WR9Yf2QPfG/MGUDHJk4vogRtXJHVxQanWHeJbon0MqOuO9yJncVI7r9WCL+UKGRPy75Ge+hbYmlPJxew"
    "moI6L3L9NWOeUyyhUzISPey5UPglB4ugPBzIcKfaAXuHRLGP6pUeK7nda8f2GHd7dqxyysicJCA0lVhh"
    "I5QLIpugM7+aKEF5EdmxBxCu7tah74VwXLNS+GmZStpcVhXOePHTxvU8LwxWP3edaEcTXORmOPJoRfe7"
    "+frOqxd8VPahJSbgLkzCiaHQ1j2rkGd2A4JF9/FUWP9xYYSLtjeoU68ZrU/1kKonVPfW/w9jQQQ3GS+q"
    "/8Hz+7g0Vr4Jy3LlezsvDP+PsHzo3mchOvvQ278AAAD//wMAUEsDBBQABgAIAAAAIQB8juxPwAIAAMcL"
    "AAARAAAAd29yZC9lbmRub3Rlcy54bWzUlttuozAQhu9X2ndA3KcGcqKoSaU221Xvqu3uA7jGCVbxQbbJ"
    "4e13zDEt3YrQq81FgLHn8/j3zMDN7ZHn3p5qw6RY+eFV4HtUEJkysVv5f34/TGLfMxaLFOdS0JV/osa/"
    "XX//dnNIqEiFtNR4gBAmOSiy8jNrVYKQIRnl2FxxRrQ0cmusiORIbreMUHSQOkVREAblndKSUGNgvXss"
    "9tj4NY4ch9FSjQ/g7IAzRDKsLT12jPBiyBxdo7gPikaAYIdR2EdNL0YtkIuqB5qNAkFUPdJ8HOmDzS3G"
    "kaI+aTmONO2T4nGkXjrxfoJLRQUMbqXm2MKj3iGO9WuhJgBW2LIXljN7AmawaDCYidcREYFXS+DT9GLC"
    "EnGZ0nyaNhS58gstktp/0vq70JPKv760HjQftiwsd43o0ebGNr56iHaV+0aSglNhS9WQpjnoKIXJmGq7"
    "Ax9Lg8Gsgew/E2DP82beQYUDS+1frW1THUMHHBJ+fXY8ryL/nBgGA07TIVqPISG8XbOJhEMGdwuPkuZM"
    "3HBg82kAUQ+wIHTgy6JhxDUDka66HYcNLKuGU52K47BO2HBgD3wfzBnApDbNLqJEja7I+WKLM2zaRHdE"
    "ellQ8xZ34mcaqd3XCuGnloXqaOxrtMeuJR7cx8kFrLqgzovcfC2Y5wwr6JScJI87ITV+ySEiKA8PMtwr"
    "T8D9Q6K4S3lLj6XdnbXneoy/7r6qvENiTwoIhiqssZXaB5PLz0lYzlPgO0vc2CMY50Ecx/fL0C+t8M6y"
    "zrqsf84VvvDSXys/CILraPlj05o2dIuL3PZHnpzpYTON74JqwSftLkZhAruFSXhrKXT1wDnkzOkfzdqH"
    "X4XbPi6s9NH6BrXuFaPZUzWkqwnlf739j5QgUlgmivJl8PxelTrGN6Isojic3YUP/4coH27vE4G6e7P+"
    "CwAA//8DAFBLAwQUAAYACAAAACEA5aDr0vwGAAD6IAAAFQAAAHdvcmQvdGhlbWUvdGhlbWUxLnhtbOxZ"
    "W48bNRR+R+I/WPOeZmaSyaVqinKltLtt1d0W8ehNnBk3nvHIdnYbISTUPvECQiqIB5CAFx4QYiWKQAjE"
    "X1h+Q6VWXH4EtmeSGScObekWVWg30saX7xx/Puf4+GTmwmt3YgIOEeOYJh3HO+c6ACVjOsFJ2HFu7o8q"
    "LQdwAZMJJDRBHWeBuPPaxVdfuQDPiwjFCEj5hJ+HHScSIj1frfKxHIb8HE1RIuemlMVQyC4LqxMGj6Te"
    "mFR9121UY4gTByQwlmpPvjz5/uTnk2NwbTrFY+RcXOofEvkvEVwNjAnbU9rRUuiL3+6eHJ/8cvLg5Pi3"
    "d2X7F/n9gZadzDz1xRe8Txg4hKTjyKUn9Ggf3REOIJALOdFxXP3nVC9eqK6EiNgiW5Ib6b9cLheYzHwt"
    "x8KDlaA79Ft1b6VfA4jYxA1b6rPSpwFwPJY7z7iUsV7QcFt+ji2BsqZFd7vp1Ux8SX9tU3+70fPrBl6D"
    "smZ9c4+j9nAQGHgNyprBBr7r+r12zcBrUNZsbODrw27THxp4DYoITmab6Eaz1Wrk6BVkSsklK7zdaLjN"
    "QQ4vUNVStGXyiXja2IvhbcpGUkA7GwqcALFI0RSOpVw3FZSDAeYpgQsHpDChXA67vufJQKy7/uqjPQDP"
    "I1iSzobGfGNI8QN8zHAqOs5lqdUpQR79+OPDuw8e3v3h4b17D+9+C3ZwGAmL3CWYhGW5P7/68K/P3gV/"
    "fPf5n/c/suN5Gf/4m/ce//TrP6kXBq2Pjx8/OH70yfu/f33fAu8yeFCG7+MYcXAVHYEbNJYbtCyADtiz"
    "SexHEJcluknIYQKVjAU9FJGBvrqABFpwPWTa8RaT6cMGfH1+2yC8F7G5wBbglSg2gLuUkh5l1j1dUWuV"
    "rTBPQvvibF7G3YDw0LZ2f83Lw3kqzwG2qexHyKB5nUiXwxAlSAA1R2cIWcTewtiw6y4eM8rpVIC3MOhB"
    "bDXJPj4woqkQuoRj6ZeFjaD0t2Gb3VugR4lN/QAdmkh5NiCxqUTEMOPrcC5gbGUMY1JG7kAR2UjuLdjY"
    "MDgX0tMhIhQMJ4hzm8w1tjDoXoEyj1ndvksWsYlkAs9syB1IaRk5oLN+BOPUyhknURn7Bp/JEIXgOhVW"
    "EtQ8Iaov/QCTre6+hZHh7ief7ZsyDdkDRM3Mme1IIGqexwWZQmRT3mWxkWK7DFujozcPjdDeQYjAIzhB"
    "CNx8w4anqWHzgvTlSGaVS8hmm8vQjFXVTxBHQBc7FsdiboTsHgrpFj67i7XEs4BJDNk2zVdnZsgMD5g8"
    "jLZ4JeOZkUoxU4fWTuIaj439bdV6PYJGWKk+t8frghn+e5ozJmVu/wsZ9MwyMrE/tW32ITEWKAJmH2Kw"
    "Y0u3UsRwfyGijpMWm1vlpuahLdxQXSt6Ypw8oQL67yofWV88+vQzC/Z0qh078HnqnG2pZL262YZbr2n6"
    "lE3wy1/SDOA8uY7kLWKBnlU0ZxXN/76i2Xaez+qYszrmrI6xi7yAOqYoXfQDoeVjH60lfupnQFNMyJ5Y"
    "ELTDdRHEZS6YjOSg7mglq0dQaSSb+fIGLmRQtwGj4k0sor0IpnJZT68Q8lx1yEFKuSyk9LBVt5og83iX"
    "TvInfKri0k89pQAUxbgbrMZl0Say0UazeES6Uq97oX4suySgZJ+FRGkxk0TNQqK5HHwCCb2zU2HRtrBo"
    "KfVbWeiv3CvysgJQPUMP6hkjGX4yxCfKT5n80run7ultxjS37Vu211ZcT8fTBolSuJkkSmEYyctkffiU"
    "fd0uXGrQU6bYpNFsvQhfq6SylhtIYvbAkTxztUCqGcO040zlDyjZjFOpj6vMBUmYdJyxyA39bzJLyrgY"
    "QB5lMD2V7T/GAjFAcCxjvewGkhTcPL+p9viSkmu7L5/l9FfZyWg6RWOxZaToyrlMiXX2OcGqQ+eS9F40"
    "OQIHZM5uQGmooOkpA04wFytrTjArBXdhxbV0lR9F491QcUQhSSOY3yjlZJ7BdXtFp7QPzXR9V2Y/38xB"
    "qJz03Lfuk4XURClpbrlA1K1pzx8v7pIvsSryvsEqS93rua69zHXbbonnvxBK1IrFDGqKsYVaMWpSO8WC"
    "oLTcKjS33RGnfRusR626IJZ1pu5tvAanB7dl5A9k9Tongmuq8lcMg/3lC8ssE+jRZXa5I8Cc4Y7ztht0"
    "630/6FfcVjCs1Gt1t9IKurVKNwhq3jDw3EHPf0caRUSxF2Rrj+SPf7LIX/Tr8Y2X/fGy9D43pnGV6nf4"
    "VS2sX/Z7vvGyP3vHD/bVvAOwtMzb/tCr+12/X+kPvEal7g8alVaz1q30/cbA78ok1Bh133HAoQZ7vcFg"
    "NAr8SqMvcXW3G1S6vVq/0mgNe/7IG9YHrgTnyfBOnj5yWywNevFvAAAA//8DAFBLAwQUAAYACAAAACEA"
    "QldEqnkEAAAMDQAAEQAAAHdvcmQvc2V0dGluZ3MueG1stFfbbts4EH1fYP/B0PMqlmTLF6FO4Uu8TRFv"
    "F3GKfaZEyiZCigJJ2XGL/fcdUqJlJ94iaZGXmJwzc2Y0HM4wHz4+cdbZEamoKCZeeBV4HVJkAtNiM/G+"
    "Piz9kddRGhUYMVGQiXcgyvt4/ftvH/aJIlqDmuoARaESnk28rdZl0u2qbEs4UleiJAWAuZAcadjKTZcj"
    "+ViVfiZ4iTRNKaP60I2CYOA1NGLiVbJIGgqf00wKJXJtTBKR5zQjzY+zkK/xW5ssRFZxUmjrsSsJgxhE"
    "oba0VI6N/ywbgFtHsvvRR+w4c3r7MHjF5+6FxEeL14RnDEopMqIUHBBnLkBatI77L4iOvq/Ad/OJlgrM"
    "w8CuTiOP30YQvSAYZOTpbRyjhqMLlqc8FL+NZ3DkoW1iw8HPBXNCoLDG2zexRC6vXWOLNNoidawiw0je"
    "FlR8pDvwNkeKvaZqauiOphLJ+k42JcOz5HZTCIlSBuFA6XTg9Ds2OvMXkmh+7JI8WbnJg3cNPeKbELyz"
    "T0oiM7go0GCiwOsaAMpT5GuNNFAkqiSM2Y6TMYLA4z7ZSMShVziJtcEkRxXTDyhda1GC0g7Bhw0dZbZF"
    "EmWayHWJMmCbi0JLwZweFn8JPYe+I+Fa1Ba5ELoQmvwtT3dgYArKD8+VGrF11n1uSwr8YvOM51zqaM4M"
    "667YrtZ1hwWTAnFI/VnXXAlMTKYqSV9fI8bAZiOMm6RddCRgIkiKyYM58rU+MLKEZK7pNzIt8OdKaQqM"
    "tnf+QgQ/CoAUxvMXKNKHQ0mWBOkKju2dnNnKWDJarqiUQt4WGGr13ZzRPCcSHFCo/RWUM5Vib/P8iSAM"
    "g/id/FaK/APK0CN6D3BNHmdCa8E/Hcot5PrXTtLVclu+8JzAyi3u4aYcVYNoHsTTJlKDtkgwjoY3i0vI"
    "/9v0ggia6CVkGPWG8+lFZNGbzceXkOk0jm6abnKOLOeD8fhibMtFbzRzGWi+myfmMWDudL0yl6fDa4s5"
    "4qmkqLMyz4Wu0Ujl44wWDk8J9GRyiqyr1IG+XwOKI8aWcIwOsAHwBFNVLkhu12yF5KblbTTkRSl01s9H"
    "LtOpifxTiqqs0b1EZX0pnErY7zeWtNB3lDu5qtK1sypgipxAVYG/7KTNU5uefaKhyGxzuUO2WK2urPz7"
    "r6a8UoqhIJH0181JZkyuTV2SFSrLurzTTTjxGN1sdWhMNOwwPDLtJt1EDRZZLKoxu0GZ+VDQbhatLHKy"
    "E72ek/VaWd/J+q0sdrK4lQ2cbGBkW2hoEqbdI9w0tzTyXDAm9gR/avEXojoJaotKsqiHIVSbqAXNdFSd"
    "XUKeYNQSTDW83UuKOYJ3VhhEA2PeaDN0EJU+0zWYUS7PGcyrpOkt3TNjW/HPYjFDOqNQnesDT9vZe1UH"
    "zqiCvlTCmNZCOuwPi4X9BIvs1rwk+rU8XkTzm2B5U8OxHe/ati4493uSz5AiuMGcaVybfo9Gi8ViGQZ+"
    "PB30/f54GPij0Sz2lzfDcLYYj2aDQf/f5s66f2Ou/wMAAP//AwBQSwMEFAAGAAgAAAAhAM45jWxrEAAA"
    "/KoAAA8AAAB3b3JkL3N0eWxlcy54bWzsXdty20YSfd+q/QcUn3YfHEkkRUmuKClJttautR3FkjfPQ2Ao"
    "IgIxXAC0rHz9zg0gyMaA6EGLUVxbrrKIS58ZzOnTmG7cfvz52yIJvvIsj0V6Pjj64XAQ8DQUUZzenw++"
    "3F2/Oh0EecHSiCUi5eeDJ54Pfv7p73/78fF1XjwlPA8kQJq/XoTng3lRLF8fHOThnC9Y/oNY8lRunIls"
    "wQq5mN0fLFj2sFq+CsViyYp4Gidx8XQwPDycDCxM1gVFzGZxyN+IcLXgaaHtDzKeSESR5vN4mZdoj13Q"
    "HkUWLTMR8jyXB71IDN6CxWkFczQGQIs4zEQuZsUP8mBsjzSUND861L8WyRrgGAcwBACTkH/DYZxajANp"
    "WceJIxzOpMKJoxqOX2dqAHlURHMUyrAc1wNlywo2Z/m8jshxnTqu4J4WaowW4ev396nI2DSRSJL1QBIX"
    "aGD1vzx+9Uf/5N/0enUIg5+kFiIRvuEztkqKXC1mN5ldtEv6z7VIizx4fM3yMI7vZAdlK4tYNvjuIs3j"
    "gdzCWV5c5DFr3DhXPxq3hHlRW30ZR/HgQLWY/yE3fmXJ+WA4LNdcqR5srEtYel+u4+mrL7f1ntRWTSXu"
    "+YBlr24vtOHR+HUS37NilckooJY0ggkWWXQlD5Z/K1YsUTsf2FEwf2tjs9xe0r1csjDWnWKzgsuYcDQ5"
    "VD1IYhWChsdn5cLnlWKKrQphG9EA5m8FewDokaFCBo5bE7/kVj77IMIHHt0WcsP5QLclV355f5PFIpMx"
    "6nxwptuUK2/5In4XRxFPazum8zjiv815+iXn0Xr9r9c6ztgVoVil8vfoZKJdJsmjt99CvlRRS25NmSLw"
    "kzJI1N6reN24Nv9vCXZkaWuyn3OmQndwtA2hu4+CGCqLvHa0zZirrWPXe6EaGu2rofG+GjreV0OTfTV0"
    "sq+GTvfVkIZ5zobiNJJnCb0/bAag7sJxqBGN4xAbGsehJTSOQypoHIcS0DgOR0fjOPwYjeNwUwROIUKX"
    "F9acfeTw9nbc3ecIP9zdpwQ/3N1nAD/c3QHfD3d3fPfD3R3O/XB3R28/3N3BGo9rplrBeymztOitspkQ"
    "RSoKHqhJb280lkosnc/S4KmTHs9IDpIAxkQ2eyLujRYyvbzbQ7RI/c/nhUoLAzELZvG9Snl6d5ynX3ki"
    "ljxgUSTxCAEzLpMyx4j4+HTGZzzjacgpHZsOVGWCQbpaTAl8c8nuybB4GhEPX4lIEhQqh5b581yJJCZw"
    "6gULM9G/a4KRxYcPcd5/rBRIcLlKEk6E9YnGxTRW/9xAw/RPDTRM/8xAw/RPDGqcUQ2RRSMaKYtGNGAW"
    "jWjcjH9SjZtFIxo3i0Y0bhat/7jdxUWiQ3x91nHUvXZ3lQh1BaJ3P27j+1RXZXsj2ZppcMMydp+x5TxQ"
    "Jexm2PoxY9u5FNFTcEdxTquQqOb12kVULTtOV/0HdAONSlwVHpG8KjwigVV4/SX2UU6T1QTtHU0+c7ua"
    "Fo2i1UidRHvLkpWZ0PZXGyv6e9haANdxlpPJoBmWwIM/qemsopMi8q172b9ja6z+stqOSqTds5AEvUxE"
    "+EATht89LXkm07KH3kjXIknEI4/oEG+LTBhfq0t+qCnpJPm3i+Wc5bHOlTYgup/qy3sXgo9s2fuAbhIW"
    "pzS8vX21YHES0M0g3t19/BDciaVKM9XA0ABeiqIQCzJMWwn8x298+k+aDl7IJDh9IjraC6LykAa7iglO"
    "MgZJRERIcpoZpzHJOVTj/Zs/TQXLIhq0m4yb24UKToR4yxZLM+kg0JaMi48y/hDMhjTef1gWq7oQlaju"
    "SMBqZcN8Nf2dh/1D3ScRkFSGflkVuv6op7ramg6u/zRhA67/FEGzKU8Pyn8JDnYDrv/BbsBRHexVwvI8"
    "dl5C9cajOtwSj/p4+yd/Fk8kIputEroBLAHJRrAEJBtCkawWaU55xBqP8IA1HvXxErqMxiMoyWm8f2Vx"
    "REaGBqNiQoNR0aDBqDjQYKQE9L9DpwbW/zadGlj/e3UMGNEUoAZG5Wekp3+iqzw1MCo/02BUfqbBqPxM"
    "g1H52ehNwGczOQmmO8XUIKl8rgZJd6JJC75YioxlT0SQbxN+zwgKpAbtJhMz9RyJSM1N3ASQqkadEE62"
    "DRwVyb/xKVnXFBZlvwgqoixJhCCqra1PONpy8961XWb6sY/eXbhJWMjnIol45jgmt63Ml2/NYxnb3dfd"
    "6FT2/BDfz4vgdl5V++swk8OdlmXCvmG2u8GmMZ+UD780mX3kUbxalB2FD1NMRt2NtUdvGI93G69nEhuW"
    "xx0tYZuT3ZbrWfKG5UlHS9jmaUdLrdMNyzY9vGHZQ6MjnLT5T5XjOZzvpM2LKuPGZtscqbJscsGTNi/a"
    "kEpwEYbqagFkp5tm3PbdxOO2x6jIjYKRkxuls67cEG0C+8y/xurMjgmaur3q7gkQ9/UkulPk/HUlTN1+"
    "44JT94e63suJU5rzoBFn1P3C1UaUcY9j53Djhugcd9wQnQOQG6JTJHKao0KSG6VzbHJDdA5Sbgh0tIJn"
    "BFy0gva4aAXtfaIVRPGJVj1mAW6IztMBNwRaqBACLdQeMwU3BEqowNxLqBAFLVQIgRYqhEALFU7AcEKF"
    "9jihQnsfoUIUH6FCFLRQIQRaqBACLVQIgRYqhEAL1XNu7zT3EipEQQsVQqCFCiHQQtXzxR5ChfY4oUJ7"
    "H6FCFB+hQhS0UCEEWqgQAi1UCIEWKoRACxVCoIQKzL2EClHQQoUQaKFCCLRQzaOG/kKF9jihQnsfoUIU"
    "H6FCFLRQIQRaqBACLVQIgRYqhEALFUKghArMvYQKUdBChRBooUIItFD1xcIeQoX2OKFCex+hQhQfoUIU"
    "tFAhBFqoEAItVAiBFiqEQAsVQqCECsy9hApR0EKFEGihQog2/7SXKF232R/hq57OO/a7X7qynfpcf5S7"
    "DjXqDlX2yo3V/VmESyEegsYHD0c63+gGEk+TWOgSteOyeh1X3xKBuvD5y1X7Ez519J4vXbLPQuhrpgB8"
    "3NUS1FTGbS5ftwRJ3rjN0+uWYNY5bou+dUtwGhy3BV2ty/KmFHk6AsZtYaZmfOQwb4vWNXM4xG0xumYI"
    "R7gtMtcM4QC3xeOa4XGggvO29XHHcZpU95cChDZ3rCGcuBHa3BJyVYZjKIyupLkRurLnRuhKoxsBxacT"
    "Bk+sGwrNsBvKj2ooMyzV/kJ1I2CphgheVAMYf6ohlDfVEMqPahgYsVRDBCzV/sHZjeBFNYDxpxpCeVMN"
    "ofyohqcyLNUQAUs1RMBS3fOE7ITxpxpCeVMNofyohpM7LNUQAUs1RMBSDRG8qAYw/lRDKG+qIZQf1SBL"
    "RlMNEbBUQwQs1RDBi2oA4081hPKmGkK1Ua2rKBtUoxiumeMmYTVD3Am5ZogLzjVDj2ypZu2ZLdUQPLMl"
    "yFXJOS5bqpPmRujKnhuhK41uBBSfThg8sW4oNMNuKD+qcdlSE9X+QnUjYKnGZUtOqnHZUivVuGyplWpc"
    "tuSmGpctNVGNy5aaqPYPzm4EL6px2VIr1bhsqZVqXLbkphqXLTVRjcuWmqjGZUtNVPc8ITth/KnGZUut"
    "VOOyJTfVuGypiWpcttRENS5baqIaly05qcZlS61U47KlVqpx2ZKbaly21EQ1LltqohqXLTVRjcuWnFTj"
    "sqVWqnHZUivVuGzpozSJCV4BdbtgWRHQvS/uHcvnBev/csIvacZzkXzlUUB7qB9QR3nwuPH5K4WtP+Qn"
    "9y/kmKk3oNceV4rMG2AtoN7xfVR9pkoZq54E9uthdrXusL1cq39nucyp7T6Hh9dvRqeX2j8PLOSOTlTN"
    "2qvIR6Dh9WevdHtTJo/3FzVOoFupemViw3rlKuX6spmrOcvM1rUTl/tYmbqPcnh1eHxh+2M/a/bA+fKT"
    "bF+vUwuSOZ7rpfUXz6bqbWNyBEbmk2f2A2inVs/CvM/pw9ekaskOpG2j9Vtz7PeWb82pjW/tOrV943Nz"
    "G5brz82p1ZfV5+ZCpf+qX9fjk4n2Gr2zjg3nA6Yjw3q1ul1FAl1eG4T1B+vKy9D1D9aZdbVPyfk4z9Dp"
    "PDY40TjPsIPzrAVr9tuQ6zO7l/2i3k73KmPGd+ZeI0t23b3Mup7uNXK6l70RhMa9Rt+Je5VD7nCvXU60"
    "D1cZ2jndxqcz9bqerjJ2uoq984fGVcYv3FVO655Shn3oKVo+9J4Sm/+vTO/6+k1Pjzh2eoS9o4vGI46/"
    "D4/QKnl5saOnD5iPwzb5gM1vaXxg8sJ9YFz3AacLaFnsNSgcn6l/2w6hvse0doe7WH3n90Lz1dMbTpze"
    "YGsVNN5w8l14QzngzxkQ9sz/qZN/Oyuh4f/0hfK/i3Etgr3qf3ii/nXh/w3FHPHMyb9lhYb/s78o/+UQ"
    "P6fi6RkP5WCz0L6y3VFhs59eqt4dpD+8tO0Lju8zOXi047+LR3e/C1XnbemzrgO3lgZNqdjpaJ09rZgm"
    "hmr5432qHO1ReUnV0+gbM1By+xVPko/M7C2W7l0TPlNykVuPDvWbOre2T81HJ5z2mb464QQ42OyMWWz3"
    "E/MZytg8NuOsxKoSfMNw62e4+o50Rx8OV7kcmlu1w3b/Nmqp2720G4OjYB1/tgJaow5cYcx6uDOEuYPS"
    "/8umaEpNhdNF6ZCIUlun63pW+v4Z7lO5RDJsiowuhkdEDNu6KD3Df1YBoM5Wn+Ihki1T53OxNSZiy5Ym"
    "Xw5b+y7gIVkxtTYXK8dErNjy4PejIXIeTL3LxcOEiAdbovtLqIO+koGkxBSdXJScEFFi62QvVBp/Ogmm"
    "8uMi4ZSIBHsW/Evo4pnz/d2UmGKMi5IzIkrsyL9QXeyrzGZenbE91mZt0xBj62saaU1YQ1HGJmyo2hko"
    "kJkrZqo4JofOFMvVwueVcjK2KkQ5xKkawhVL7Kv8zci9gHs71kekj/pVOSwPPKvGfj2XLtcc2/NtfXZt"
    "1tGJcs1go5f0VWPN1dzO8TKz2v1z1qzh6ive2wRVGyiUXIK1itmyghJzulqYH3ECb7uyG5+5xI2dhQDu"
    "j2wCst/Ed4MSF/l9BbrpRG7OX/jE8Zkpa1am+drANjNmLYUmNVKbIId2MuN5dq3fzab3+D0sLVXuynW7"
    "QJstc8vxofrXhTXqNHg9VI109FVJjVM3CzslsteRa3ZZddVk/cGO7bHSjzusN+/yYTgUI1s/QzlkrK9w"
    "qetT6uV71hXb5nId3aU6aPtGuuo1eduHDd6jh3OUBo9AnSh3e8ceb9OyY9Ec2jY/s7LLPbqEuHpzbZFu"
    "5JNHLC8j/ddcF9X75dKT7Ae7/1C32Kkf0r9UPNHq08PuWRavrqA+c0tKBvbIdj1XoZaMW9U0djrRvdEX"
    "dM2S3qVv8P9Ty6DAj1pdt+/pYEMkOzz2xem+NUauX9rpGsD1Hn2jZHmpDxUlp6ZVO1q5DCrJFVvSjB2Y"
    "RJb3X3pFUlWI4tAR1Q1ORlytwXO7pMWztvBY3vDQcC+J6+G0StkFm+rH5OTfrQAgF5ciV499li/rrO2j"
    "w0e1y9nkVHfhYI3XcJtTsLOIgzzF76g22kHbpsBsoig0lkS6GXENP9abroVQx7x9KDOzGuNNBun/3oTy"
    "ptqgbVNgNvX1Jstvb28qf+U//Q8AAP//AwBQSwMEFAAGAAgAAAAhAO8KKU5OAQAAfgMAABQAAAB3b3Jk"
    "L3dlYlNldHRpbmdzLnhtbJzTX2vCMBAA8PfBvkPJu6bKFClWYQzHXsZg2weI6dWGJbmSi6vu0+/aqXP4"
    "YveS//fjLiHz5c7Z5BMCGfS5GA1TkYDXWBi/ycX722owEwlF5Qtl0UMu9kBiubi9mTdZA+tXiJFPUsKK"
    "p8zpXFQx1pmUpCtwioZYg+fNEoNTkadhI50KH9t6oNHVKpq1sSbu5ThNp+LAhGsULEuj4QH11oGPXbwM"
    "YFlET5Wp6ag112gNhqIOqIGI63H2x3PK+BMzuruAnNEBCcs45GIOGXUUh4/SbuTsLzDpB4wvgKmGXT9j"
    "djAkR547pujnTE+OKc6c/yVzBlARi6qXMj7eq2xjVVSVoupchH5JTU7c3rV35HT2tPEY1NqyxK+e8MMl"
    "Hdy2XH/bdUPYdettCWLBHwLraJz5ghWG+4ANQZDtsrIWm5fnR57IP79m8Q0AAP//AwBQSwMEFAAGAAgA"
    "AAAhAFU6fTT3AQAALgcAABIAAAB3b3JkL2ZvbnRUYWJsZS54bWzck99umzAUxu8n7R0s3zcYEtIUlVTr"
    "2ki72cXUPYBjTLDmP8jHCcnbzzaEpoqmlUmTpoEEh+9wfjrnfHD/cFQSHbgFYXSJ0xnBiGtmKqF3Jf7+"
    "srlZYQSO6opKo3mJTxzww/rjh/uuqI12gHy9hkKxEjfOtUWSAGu4ojAzLdc+WRurqPOPdpcoan/s2xtm"
    "VEud2Aop3CnJCFniAWPfQzF1LRh/MmyvuHaxPrFceqLR0IgWzrTuPbTO2Kq1hnEAP7OSPU9RoUdMurgC"
    "KcGsAVO7mR9m6CiifHlKYqTkKyCfBsiuAEvGj9MYq4GR+MpLjqimcZYjR1QXnD9r5gIAlauaSZTsvNck"
    "1FJHGwrNJZFPayofcScVdqRY8WWnjaVb6UnedeSNQxEcrn7+cIshP0Y9jIDXw6+AukJT5Ss/tc5AlFlD"
    "LfCQOVBZYkJwEt+mSsjTWYVOAPSJVjjWnPUDtSK00qdA7HxiD1tSYv+FEpKtbnGvpIEcj/mgZKNCBmX+"
    "VmGREx/Tu82gvHJin0k/1vV4vi0Zx2upNsDTcbyMPJIlWZBgVX8u+ub/0h6ewxqeN7H/fg+fvXK7yh+v"
    "9nD3+z30nAl7CDajJwGtpKf/1+4XoTigr7xD34yi+hfGZ974Ocm9+bmP55OMt5H7Dxs/BLD+CQAA//8D"
    "AFBLAwQUAAYACAAAACEA+ifkVXwBAAARAwAAEQAIAWRvY1Byb3BzL2NvcmUueG1sIKIEASigAAEAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAjJLfT4MwEIDfTfwfSN+hBXUawtCo8UkTE2c0vtX2xuroj7TdkP/eAhuTaKJvPe67"
    "r8ddi8tPWUdbsE5oNUdpQlAEimkuVDVHz4u7+AJFzlPFaa0VzFELDl2Wx0cFMznTFh6tNmC9ABcFk3I5"
    "M3O08t7kGDu2AkldEggVkkttJfUhtBU2lK1pBTgjZIYleMqpp7gTxmY0op2Ss1FpNrbuBZxhqEGC8g6n"
    "SYoPrAcr3a8FfeYbKYVvDfyK7pMj/enECDZNkzQnPRr6T/Hrw/1T/6uxUN2sGKCy4Cz3wtdQFvhwDCe3"
    "ef8A5ofPYxDOzAL12pZqY1ta8baFbSJhTRU5v6okFXXCtOyr9mS3gzW0jbbcBd8kChgHx6wwPmx2uG3y"
    "IdA1df4hrHopgF+3f138s6BzWNiK7u2UWU+MYbFbxNAs8CgMMB/Gvc+8nNzcLu5QmZFsFpPTOM0W5CLP"
    "znJC3rp+J/UHodw18H/jbGrcC4aRTR9x+QUAAP//AwBQSwMEFAAGAAgAAAAhAPU2xKVyAQAAxgIAABAA"
    "CAFkb2NQcm9wcy9hcHAueG1sIKIEASigAAEAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
    "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAnFLLTsMwELwj8Q9R7q3TihZU"
    "bV2hVogDj0pN27PlbBILx7Zsg9q/Z9NACOJGTjuzu6PZiWF1anTygT4oa5bpZJylCRppC2WqZbrPH0Z3"
    "aRKiMIXQ1uAyPWNIV/z6CrbeOvRRYUhIwoRlWsfoFowFWWMjwpjahjql9Y2IBH3FbFkqiRsr3xs0kU2z"
    "bM7wFNEUWIxcL5h2iouP+F/RwsrWXzjkZ0d6HHJsnBYR+Uu7qYH1BOQ2Cp2rBvmE6B7AVlQYWq4r4Gh9"
    "QXgyBdaVsK6FFzJSeHx+Q4MDDPfOaSVFpFj5s5LeBlvG5PXiNWn3gQ1HgPzvUL57Fc88AzaE8KQMGZgB"
    "6wpy5kXlhau/7PUIdlJoXNPlvBQ6ILAfAta2ccKQHOsr0nsLe5fbTZvE18pvcnDkUcV654QkC7ez6fDc"
    "QQd2xGJB/nsLPQGP9De8bvVp11RYfM/8bbQBHrpnySfzcUbfJbFvju7u3wv/BAAA//8DAFBLAQItABQA"
    "BgAIAAAAIQBKvAJxbQEAACgGAAATAAAAAAAAAAAAAAAAAAAAAABbQ29udGVudF9UeXBlc10ueG1sUEsB"
    "Ai0AFAAGAAgAAAAhAB6RGrfvAAAATgIAAAsAAAAAAAAAAAAAAAAApgMAAF9yZWxzLy5yZWxzUEsBAi0A"
    "FAAGAAgAAAAhAJWp/5nWCwAAhYcAABEAAAAAAAAAAAAAAAAAxgYAAHdvcmQvZG9jdW1lbnQueG1sUEsB"
    "Ai0AFAAGAAgAAAAhABEXoNkNAQAAOQQAABwAAAAAAAAAAAAAAAAAyxIAAHdvcmQvX3JlbHMvZG9jdW1l"
    "bnQueG1sLnJlbHNQSwECLQAUAAYACAAAACEArshHZ8ICAADNCwAAEgAAAAAAAAAAAAAAAAAaFQAAd29y"
    "ZC9mb290bm90ZXMueG1sUEsBAi0AFAAGAAgAAAAhAHyO7E/AAgAAxwsAABEAAAAAAAAAAAAAAAAADBgA"
    "AHdvcmQvZW5kbm90ZXMueG1sUEsBAi0AFAAGAAgAAAAhAOWg69L8BgAA+iAAABUAAAAAAAAAAAAAAAAA"
    "+xoAAHdvcmQvdGhlbWUvdGhlbWUxLnhtbFBLAQItABQABgAIAAAAIQBCV0SqeQQAAAwNAAARAAAAAAAA"
    "AAAAAAAAACoiAAB3b3JkL3NldHRpbmdzLnhtbFBLAQItABQABgAIAAAAIQDOOY1saxAAAPyqAAAPAAAA"
    "AAAAAAAAAAAAANImAAB3b3JkL3N0eWxlcy54bWxQSwECLQAUAAYACAAAACEA7wopTk4BAAB+AwAAFAAA"
    "AAAAAAAAAAAAAABqNwAAd29yZC93ZWJTZXR0aW5ncy54bWxQSwECLQAUAAYACAAAACEAVTp9NPcBAAAu"
    "BwAAEgAAAAAAAAAAAAAAAADqOAAAd29yZC9mb250VGFibGUueG1sUEsBAi0AFAAGAAgAAAAhAPon5FV8"
    "AQAAEQMAABEAAAAAAAAAAAAAAAAAETsAAGRvY1Byb3BzL2NvcmUueG1sUEsBAi0AFAAGAAgAAAAhAPU2"
    "xKVyAQAAxgIAABAAAAAAAAAAAAAAAAAAxD0AAGRvY1Byb3BzL2FwcC54bWxQSwUGAAAAAA0ADQBAAwAA"
    "bEAAAAAA"
)


# ✅ DÜZEDIŞ #2: Şablon ÝÜKLENIŞI — iki çeşmeli:
# 1-nji: Eger "template.docx" faýly bar bolsa — şondan oka (iň ygtybarly)
# 2-nji: Ýogsa — base64-den dekompressiýa et
_TEMPLATE_BYTES_CACHE = None

# Birinji: template.docx faýlyndan synanyş
_TEMPLATE_FILE_PATHS = ["template.docx", "/app/template.docx", "./template.docx"]
for _path in _TEMPLATE_FILE_PATHS:
    try:
        if os.path.exists(_path):
            with open(_path, "rb") as _f:
                _candidate = _f.read()
            if _candidate[:4] == b'PK\x03\x04':
                _TEMPLATE_BYTES_CACHE = _candidate
                log.info(f"✅ Şablon faýldan ýüklendi: {_path} ({len(_candidate)} baýt)")
                break
            else:
                log.warning(f"⚠️ {_path} ZIP däl, indiki çeşme synanyşýar")
    except Exception as e:
        log.warning(f"⚠️ {_path} okalmady: {e}")

# Ikinji: Base64-den synanyş (faýl tapylmasa)
if _TEMPLATE_BYTES_CACHE is None:
    try:
        # Başga-başga ýerlerde WHITE-SPACE bozulýan bolsa - arassala
        _clean_b64 = "".join(TEMPLATE_B64.split())
        _TEMPLATE_BYTES_CACHE = base64.b64decode(_clean_b64)
        if _TEMPLATE_BYTES_CACHE[:4] != b'PK\x03\x04':
            log.error(f"❌ TEMPLATE_B64 ZIP faýly däl! Başlangyç baýtlar: {_TEMPLATE_BYTES_CACHE[:10]}")
            _TEMPLATE_BYTES_CACHE = None
        else:
            # Hakykatdanam açylyp bilýärmi barla
            try:
                import zipfile
                with zipfile.ZipFile(io.BytesIO(_TEMPLATE_BYTES_CACHE)) as _zf:
                    _names = _zf.namelist()
                    if "word/document.xml" not in _names:
                        log.error(f"❌ Şablonda word/document.xml ýok! Faýllar: {_names[:5]}")
                        _TEMPLATE_BYTES_CACHE = None
                    else:
                        log.info(f"✅ Şablon base64-den ýüklendi ({len(_TEMPLATE_BYTES_CACHE)} baýt, {len(_names)} içki faýl)")
            except Exception as zex:
                log.error(f"❌ Şablon ZIP hökmünde açylmaýar: {zex}")
                _TEMPLATE_BYTES_CACHE = None
    except Exception as e:
        log.error(f"❌ TEMPLATE_B64 açylmady: {e}")
        _TEMPLATE_BYTES_CACHE = None

if _TEMPLATE_BYTES_CACHE is None:
    log.error("🔴 ŞABLON ÝÜKLENMEDI — Referat/Doklad işlemez! template.docx faýlyny repositora goş.")


def _get_template_bytes() -> bytes:
    if _TEMPLATE_BYTES_CACHE is None:
        raise RuntimeError(
            "Şablon ýüklenmedi! template.docx faýly ýok ýa-da TEMPLATE_B64 bozuk. "
            "Çözgüt: template.docx faýlyny repositoryň köküne goşuň."
        )
    # Her çagyryşda täze kopiýa gaýtarýarys — io.BytesIO sebäpli bozulmasyn
    return bytes(_TEMPLATE_BYTES_CACHE)


class St(StatesGroup):
    s01  = State()
    s02  = State()
    s02b = State()
    s03  = State()
    s04  = State()
    s05  = State()
    s06  = State()
    s07  = State()
    s08  = State()
    s09  = State()
    s09b = State()
    s10  = State()
    s11  = State()
    s11b = State()
    s12  = State()
    s13  = State()
    szad1 = State()
    szad2 = State()
    szad3 = State()
    s02b_more = State()


PENDING: dict[int, dict] = {}
SEEN_USERS: set = set()
REQ_ITEMS: dict[int, list] = {}
PAYMENT_PENDING: dict[int, dict] = {}
ACTIVE_GENERATES: set = set()
CANCELLED_GENERATES: set = set()
ZADANIYE_GENERATING: set = set()
SVC_RU = {"referat": "Реферат", "doklad": "Доклад", "zadaniye": "Задание"}
SVC_TM = {"referat": "Referat 📄", "doklad": "Doklad 🎤", "zadaniye": "Zadaniye 📝"}


def kb(*rows) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=t, callback_data=c) for t, c in row]
        for row in rows
    ])

KB_SVC  = kb(
    [("📄  Referat  — 300 ₽","svc:referat")],
    [("🎤  Doklad   — 300 ₽","svc:doklad")],
    [("📝  Zadaniye — 150 ₽","svc:zadaniye")])
KB_REQ  = kb([("✅  Talaplы — öz talaplarym bar","req:yes")],[("➡️  Talapsyz — adaty GOST","req:no")],
              [("🔙  Yza — başa gaýt","back:start")])
KB_CRS  = kb([("1-nji kurs","crs:1"),("2-nji kurs","crs:2"),("3-nji kurs","crs:3")],
              [("4-nji kurs","crs:4"),("5-nji kurs","crs:5")],
              [("🔙  Yza — başa gaýt","back:start")])
KB_SEC  = kb([("2 bölüm","sec:2"),("3 bölüm","sec:3")],
              [("🔙  Yza — başa gaýt","back:start")])
KB_SPC  = kb([("📋  Adaty 1.5 — dowam et","spc:default")],[("✏️  Üýtgetmek","spc:custom")],
              [("🔙  Yza — başa gaýt","back:start")])
KB_SPCV = kb([("1.0","spv:1.0"),("1.25","spv:1.25"),("1.5","spv:1.5")])
KB_SKIP = kb([("⏭️  Mugallym ýok — geç","skip:teacher")],
              [("🔙  Yza — başa gaýt","back:start")])
KB_BACK     = kb([("🔙  Yza","back:start")])
KB_REQ_DONE = kb([("✅  Talaplarym taýar — dowam et","req:done")],
                 [("🔙  Yza — başa gaýt","back:start")])

def kb_src() -> InlineKeyboardMarkup:
    rows, row = [], []
    for i in range(8, 21):
        row.append((str(i), f"src:{i}"))
        if len(row) == 5:
            rows.append(row); row = []
    if row: rows.append(row)
    return kb(*rows)


async def ask(obj, text: str, markup=None) -> Message:
    kw = dict(parse_mode="HTML", reply_markup=markup)
    if isinstance(obj, CallbackQuery):
        try:
            await obj.message.edit_text(text, **kw); return obj.message
        except Exception:
            return await obj.message.answer(text, **kw)
    return await obj.answer(text, **kw)

def md_clean(t: str) -> str:
    t = re.sub(r"\*\*(.*?)\*\*", r"\1", t)
    t = re.sub(r"\*(.*?)\*", r"\1", t)
    t = re.sub(r"^#{1,6}\s*", "", t, flags=re.M)
    return t.strip()

def spc_float(k: str) -> float:
    return {"default":1.5,"1.0":1.0,"1.25":1.25,"1.5":1.5}.get(k, 1.5)

def spc_str(k: str) -> str:
    return {"default":"1.5","1.0":"1.0","1.25":"1.25","1.5":"1.5"}.get(k, "1.5")

def lv(spc) -> int:
    return {1.0:240, 1.25:300, 1.5:360}.get(spc, 360)


def build_zadaniye_prompt(d: dict) -> str:
    req_text = d.get("req_text", "").strip()
    if req_text:
        extra_block = (
            f"═══ СТУДЕНТИҢ ÝÖRITE TALAPLARY (HÖKMAN ÝERINE ÝETIR) ═══\n"
            f"{req_text}\n"
            f"Ýokardaky talaplary doly ýerine ýetir!\n"
            f"═══════════════════════════════════════════════════════════\n\n"
        )
    else:
        extra_block = ""
    return (
        f"{extra_block}"
        f"Напиши академическое задание строго на РУССКОМ языке по теме: «{d['theme']}»\n\n"
        f"Требования к тексту:\n"
        f"• Объём: 2–3 страницы A4 (~3600–4500 символов)\n"
        f"• Шрифт 14 пт, интервал 1.5, абзацный отступ 1.5 см\n"
        f"• Стиль: академический, связный, без заголовков\n"
        f"• Структура: вводное предложение → развёрнутые абзацы → вывод\n"
        f"• Все перечисления нумеруй: 1. 2. 3. (не используй маркеры •/—)\n\n"
        f"Начинай СРАЗУ с первого предложения абзаца, без каких-либо предисловий."
    )


def build_prompt(d: dict) -> str:
    svc   = SVC_RU.get(d["service"], "Реферат")
    pages = int(d["pages"])
    secs  = int(d["sections"])
    spc   = spc_str(d.get("spacing", "default"))
    cpp   = 1500
    total = pages * cpp
    ic    = int(total * 0.15)
    cc    = int(total * 0.65 / secs)
    nc    = int(total * 0.10)
    chs   = ""
    for i in range(1, secs + 1):
        chs += (f"\n##ГЛАВА_{i}##\n"
                f"Первая строка — заголовок: «{i}. [Название главы]»\n"
                f"Текст: ~{cc} символов, сплошными абзацами. БЕЗ подразделов.\n")
    teacher = d.get("teacher","______________")
    tp      = d.get("teacher_position","")
    tline   = f"{teacher}" if not tp else f"{tp} {teacher}"
    req_text = d.get("req_text", "").strip()
    if req_text:
        extra_block = (
            f"═══ СТУДЕНТИҢ ÝÖRITE TALAPLARY (HÖKMAN ÝERINE ÝETIR) ═══\n"
            f"{req_text}\n"
            f"Ýokardaky talaplary doly ýerine ýetir — olary äsgermezlik etme!\n"
            f"═══════════════════════════════════════════════════════════\n\n"
        )
    else:
        extra_block = ""
    return (
        f"{extra_block}"
        f"Ты академический автор. Напиши {svc} строго на РУССКОМ языке.\n\n"
        f"ПАРАМЕТРЫ:\n"
        f"Тип: {svc} | Дисциплина: {d['subject']} | Тема: {d['theme']}\n"
        f"Университет: {d['university']}\n"
        f"Студент: {d['fullname']}, {d['course']} курс, гр. {d['group']}\n"
        f"Преподаватель: {tline}\n"
        f"Глав: {secs} | Страниц: {pages} | Источников: {d['sources']} | Интервал: {spc}\n\n"
        f"ОБЪЁМ: РОВНО {pages} стр. A4 = ~{total} символов.\n"
        f"• Введение: ~{ic} символов\n• Каждая глава: ~{cc} символов\n• Заключение: ~{nc} символов\n\n"
        f"СТРУКТУРА — строго эти маркеры:\n\n"
        f"##ВВЕДЕНИЕ##\n~{ic} символов. Актуальность, цель, {secs} задачи, методология.\n"
        f"{chs}\n"
        f"##ЗАКЛЮЧЕНИЕ##\n~{nc} символов. Выводы, итоги, значение.\n\n"
        f"##СПИСОК_ЛИТЕРАТУРЫ##\n"
        f"Ровно {d['sources']} источников, ГОСТ Р 7.0.5-2008.\n"
        f"Формат: «N. Автор. Название. — Город : Издательство, Год. — Стр. с.»\n"
        f"Только 1. 2. 3. — НЕТ 1.1/1.2\n\n"
        f"ПРАВИЛА:\n"
        f"• только русский язык\n"
        f"• НЕТ подразделов 1.1/1.2\n"
        f"• Все перечисления нумеруй: 1. 2. 3. (не используй маркеры •/—)\n"
        f"• Начинай с ##ВВЕДЕНИЕ##:"
    )


STAGES = [
    (5,  "🔍 Tema seljerilýär..."),
    (15, "📚 Çeşmeler gözlenýär..."),
    (30, "📖 Giriş ýazylýar..."),
    (50, "📝 Esasy bölümler..."),
    (70, "📝 Dowam edýär..."),
    (85, "🔚 Netije ýazylýar..."),
    (93, "📑 Çeşmeler sanawy..."),
    (97, "🔧 Jemlenýär..."),
]

async def call_deepseek(d: dict, on_progress) -> str:
    prompt  = build_zadaniye_prompt(d) if d.get("service") == "zadaniye" else build_prompt(d)
    # ✅ DÜZEDIŞ #1: Accept-Encoding: identity → "invalid distance too far back" ýalňyşlygyny aýyrýar
    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type":  "application/json",
        "Accept-Encoding": "identity",   # gzip/br gysyşy öçür — bozulan paketler sebäpli ýalňyşlyk çykmaz
        "Accept": "application/json",
    }
    req_items = d.get("req_items", [])
    extra_texts = []
    for item in req_items:
        if item["type"] == "text":
            extra_texts.append(item["content"])
        elif item["type"] == "image":
            extra_texts.append("[Ulanyjy surat iberdi — suradyň ýanyndaky tekst talabyna görä hereket et]")
    if extra_texts:
        user_content_final = "\n\n".join(extra_texts) + "\n\n" + prompt
    else:
        user_content_final = prompt
    system_prompt = (
        "Ты профессиональный академический автор. "
        "Пиши ТОЛЬКО на русском языке. "
        "Если в начале задания есть раздел «СТУДЕНТИҢ ÝÖRITE ТАЛАРЫ» — "
        "это особые требования студента, их ОБЯЗАТЕЛЬНО нужно выполнить полностью. "
        "Все перечисления нумеруй: 1. 2. 3. — никаких маркеров •/—. "
        "Без вступлений и пояснений."
    )
    body = {
        "model":       DEEPSEEK_MODEL,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_content_final},
        ],
        "max_tokens":  6000,
        "temperature": 0.7,
        "stream":      False,
    }
    result: dict = {}
    error:  dict = {}
    done:   dict = {"flag": False}

    async def fetch():
        max_retries = 5
        last_exc    = None
        for attempt in range(1, max_retries + 1):
            try:
                async with httpx.AsyncClient(
                    timeout=httpx.Timeout(connect=60.0, read=600.0, write=60.0, pool=30.0),
                    limits=httpx.Limits(max_keepalive_connections=5, max_connections=10),
                    http2=False,
                ) as cl:
                    r = await cl.post(DEEPSEEK_URL, headers=headers, json=body)
                    if r.status_code != 200:
                        raise RuntimeError(f"HTTP {r.status_code}: {r.text[:300]}")
                    resp = r.json()
                    text = resp["choices"][0]["message"]["content"]
                    if not text or not text.strip():
                        raise RuntimeError("DeepSeek boş jogap iberdi")
                    result["text"] = text.strip()
                    return
            except (httpx.ConnectError, httpx.ConnectTimeout,
                    httpx.ReadTimeout, httpx.RemoteProtocolError) as e:
                last_exc = e
                wait = min(attempt * 10, 60)
                log.warning(f"Bağlantı {attempt}/{max_retries}: {type(e).__name__} — {wait}s garaşylýar")
                await asyncio.sleep(wait)
            except RuntimeError as e:
                error["exc"] = e
                done["flag"] = True
                return
            except Exception as e:
                last_exc = e
                wait = min(attempt * 10, 60)
                log.warning(f"Ýalňyşlyk {attempt}/{max_retries}: {e} — {wait}s garaşylýar")
                await asyncio.sleep(wait)
        error["exc"] = last_exc
        done["flag"] = True

    async def ticker():
        stage_idx = 0
        elapsed   = 0
        while not done["flag"]:
            await asyncio.sleep(6)
            elapsed += 6
            if stage_idx < len(STAGES):
                pct, status = STAGES[stage_idx]
                stage_idx += 1
            else:
                pct    = 97
                mins   = elapsed // 60
                secs_e = elapsed % 60
                status = f"⏳ Taýarlanýar... {mins}:{secs_e:02d}"
            try:
                await on_progress(pct, status)
            except Exception:
                pass

    done["flag"] = False
    ft = asyncio.create_task(fetch())
    tt = asyncio.create_task(ticker())
    try:
        await ft
    finally:
        done["flag"] = True
        tt.cancel()
        try:
            await tt
        except asyncio.CancelledError:
            pass
    if "exc" in error:
        raise error["exc"]
    await on_progress(100, "✅ Taýar!")
    return result["text"]


def parse_ai(raw: str, secs: int) -> dict:
    def _between(text, start, *ends):
        s = text.find(start)
        if s == -1: return ""
        s += len(start); best = len(text)
        for e in ends:
            p = text.find(e, s)
            if p != -1 and p < best: best = p
        return text[s:best].strip()

    intro_raw = _between(raw, "##ВВЕДЕНИЕ##", "##ГЛАВА_1##", "##ЗАКЛЮЧЕНИЕ##", "##СПИСОК_ЛИТЕРАТУРЫ##")
    chapters: list = []
    for i in range(1, secs + 1):
        nxt    = f"##ГЛАВА_{i+1}##" if i < secs else "##ЗАКЛЮЧЕНИЕ##"
        ch_raw = _between(raw, f"##ГЛАВА_{i}##", nxt, "##СПИСОК_ЛИТЕРАТУРЫ##")
        if not ch_raw: continue
        lines = [ln.strip() for ln in ch_raw.splitlines() if ln.strip()]
        if lines and re.match(r"^\d+\.", lines[0]):
            title = md_clean(lines[0]); body = lines[1:]
        else:
            title = f"{i}. Глава {i}"; body = lines
        chapters.append({"title": title, "lines": body})
    conc_raw = _between(raw, "##ЗАКЛЮЧЕНИЕ##", "##СПИСОК_ЛИТЕРАТУРЫ##")
    src_raw  = _between(raw, "##СПИСОК_ЛИТЕРАТУРЫ##")
    raw_srcs = [ln.strip() for ln in src_raw.splitlines() if ln.strip() and not ln.strip().startswith("##")]
    sources: list = []
    for ln in raw_srcs:
        ln = re.sub(r"^(\d+)\.\d+\.", r"\1.", ln)
        if re.match(r"^\d+\.", ln): sources.append(ln)
        elif sources: sources[-1] += " " + ln
    return dict(
        intro      = [ln.strip() for ln in intro_raw.splitlines() if ln.strip()],
        chapters   = chapters,
        conclusion = [ln.strip() for ln in conc_raw.splitlines() if ln.strip()],
        sources    = sources,
    )


def _sf(run, size_pt=14, bold=False, italic=False):
    run.font.name = "Times New Roman"; run.font.size = Pt(size_pt)
    run.bold = bold; run.italic = italic
    rpr = run._r.get_or_add_rPr()
    rf  = OxmlElement("w:rFonts")
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rf.set(qn(a), "Times New Roman")
    old = rpr.find(qn("w:rFonts"))
    if old is not None: rpr.remove(old)
    rpr.insert(0, rf)
    for tag in ("w:sz","w:szCs"):
        el = rpr.find(qn(tag))
        if el is None: el = OxmlElement(tag); rpr.append(el)
        el.set(qn("w:val"), str(size_pt * 2))

def _spf(para, align, line, first_line=0, left=0, hanging=0, space_after=0):
    pPr = para._p.get_or_add_pPr()
    sp  = pPr.find(qn("w:spacing"))
    if sp is None: sp = OxmlElement("w:spacing"); pPr.append(sp)
    sp.set(qn("w:after"), str(space_after))
    sp.set(qn("w:line"), str(line))
    sp.set(qn("w:lineRule"), "auto")
    ind_old = pPr.find(qn("w:ind"))
    if ind_old is not None: pPr.remove(ind_old)
    if first_line or left or hanging:
        ind = OxmlElement("w:ind")
        if first_line: ind.set(qn("w:firstLine"), str(first_line))
        if left:       ind.set(qn("w:left"), str(left))
        if hanging:    ind.set(qn("w:hanging"), str(hanging))
        pPr.append(ind)
    jc = pPr.find(qn("w:jc"))
    if jc is None: jc = OxmlElement("w:jc"); pPr.append(jc)
    jc.set(qn("w:val"), align)

def _para(doc, text, *, bold=False, italic=False, center=False,
          size_pt=14, line=360, first_line=851, space_after=0):
    p = doc.add_paragraph(); r = p.add_run(text)
    _sf(r, size_pt, bold, italic)
    _spf(p, "center" if center else "both", line, first_line=first_line, space_after=space_after)

def _page_break(doc):
    p = doc.add_paragraph(); r = p.add_run()
    br = OxmlElement("w:br"); br.set(qn("w:type"), "page"); r._r.append(br)
    _spf(p, "both", 360)

def _init_heading(doc):
    try:    st = doc.styles["Heading 1"]
    except: st = doc.styles.add_style("Heading 1", 1)
    st.font.name = "Times New Roman"; st.font.size = Pt(14)
    st.font.bold = True; st.font.color.rgb = RGBColor(0, 0, 0)
    pf = st.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    el = st.element; pPr = el.get_or_add_pPr()
    ol = pPr.find(qn("w:outlineLvl"))
    if ol is None: ol = OxmlElement("w:outlineLvl"); pPr.append(ol)
    ol.set(qn("w:val"), "0")

def _heading(doc, text: str, line_v: int):
    p = doc.add_paragraph(style="Heading 1")
    for run in p.runs: run.text = ""
    r = p.add_run(text); _sf(r, 14, True)
    r.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before      = Pt(0)
    p.paragraph_format.space_after       = Pt(0)
    _spf(p, "center", line_v, first_line=0)

def _set_p_text(new_p, text: str):
    r_els = new_p.findall(f".//{qn('w:r')}")
    written = False
    for r_el in r_els:
        for t_el in r_el.findall(qn("w:t")):
            if not written:
                t_el.text = text
                t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                written = True
            else:
                t_el.text = ""

def _copy_template_title(doc: Document, d: dict):
    svc_ru      = SVC_RU.get(d["service"], "Реферат")
    teacher     = d.get("teacher", "").strip()
    t_position  = d.get("teacher_position", "").strip()
    has_teacher = bool(teacher and teacher != "______________")
    subject     = d.get("subject", "").strip()
    university  = d.get("university", "").strip()
    try:
        tmpl_bytes = _get_template_bytes()
        tmpl = Document(io.BytesIO(tmpl_bytes))
    except Exception as e:
        log.error(f"❌ Şablon açylmady: {e}")
        raise RuntimeError(f"Şablon açylmady: {e}")
    body   = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    skip_teacher = set(range(25, 34)) if not has_teacher else set()
    skip = skip_teacher | {37, 38, 39}

    def _append(p_el):
        if sectPr is not None:
            body.insert(list(body).index(sectPr), p_el)
        else:
            body.append(p_el)

    for i, tp in enumerate(tmpl.paragraphs):
        if i in skip:
            continue
        new_p = copy.deepcopy(tp._element)
        if i == 3:
            _set_p_text(new_p, f"«{university}»")
        elif i == 4:
            _set_p_text(new_p, "")
        elif i == 9:
            if subject:
                new_text = (f"{svc_ru} по дисциплине {subject}\n"
                            f"на тему: «{d['theme']}»")
            else:
                new_text = f"{svc_ru} на тему: «{d['theme']}»"
            _set_p_text(new_p, new_text)
        elif i == 17:
            _set_p_text(new_p, f"Студент {d['course']} курса, группы {d['group']}")
        elif i == 18:
            _set_p_text(new_p, d["fullname"])
        elif i == 25 and has_teacher:
            проверил = f"Проверил: {t_position}" if t_position else "Проверил:"
            _set_p_text(new_p, проверил)
        elif i == 26 and has_teacher:
            _set_p_text(new_p, teacher)
            _append(new_p)
            p19 = copy.deepcopy(tmpl.paragraphs[19]._element)
            _append(p19)
            continue
        _append(new_p)

    p39 = copy.deepcopy(tmpl.paragraphs[39]._element)
    _set_p_text(p39, "2026г.")
    _append(p39)

def _auto_toc(doc, line_v: int):
    _para(doc, "Содержание", bold=True, center=True, size_pt=14,
          line=line_v, first_line=0, space_after=0)
    p = doc.add_paragraph()
    r1 = p.add_run(); fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "begin"); r1._r.append(fc)
    r2 = p.add_run(); ins = OxmlElement("w:instrText")
    ins.set(qn("xml:space"), "preserve")
    ins.text = ' TOC \\o "1-1" \\h \\z \\u '
    r2._r.append(ins)
    r3 = p.add_run(); fs = OxmlElement("w:fldChar")
    fs.set(qn("w:fldCharType"), "separate"); r3._r.append(fs)
    r4 = p.add_run(); fe = OxmlElement("w:fldChar")
    fe.set(qn("w:fldCharType"), "end"); r4._r.append(fe)
    settings = doc.settings.element
    uf = OxmlElement("w:updateFields")
    uf.set(qn("w:val"), "true")
    settings.append(uf)

def _add_page_numbers(doc: Document) -> None:
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI
    FOOTER_XML = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:ftr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        '       xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<w:p><w:pPr><w:jc w:val="center"/><w:spacing w:before="0" w:after="0"/></w:pPr>'
        '<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>'
        '<w:sz w:val="28"/><w:szCs w:val="28"/></w:rPr><w:fldChar w:fldCharType="begin"/></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>'
        '<w:sz w:val="28"/><w:szCs w:val="28"/></w:rPr><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>'
        '<w:sz w:val="28"/><w:szCs w:val="28"/></w:rPr><w:fldChar w:fldCharType="separate"/></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>'
        '<w:sz w:val="28"/><w:szCs w:val="28"/></w:rPr><w:fldChar w:fldCharType="end"/></w:r>'
        '</w:p></w:ftr>'
    )
    FOOTER_EMPTY_XML = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:ftr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:p><w:pPr><w:jc w:val="center"/></w:pPr></w:p></w:ftr>'
    )
    FOOTER_CT  = 'application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml'
    FOOTER_REL = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer'
    footer_part = Part(PackURI('/word/footer_pnum.xml'), FOOTER_CT,
                       FOOTER_XML.encode('utf-8'), doc.part.package)
    rId = doc.part.relate_to(footer_part, FOOTER_REL)
    footer_first_part = Part(PackURI('/word/footer_first.xml'), FOOTER_CT,
                             FOOTER_EMPTY_XML.encode('utf-8'), doc.part.package)
    rId_first = doc.part.relate_to(footer_first_part, FOOTER_REL)
    body   = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is None:
        sectPr = OxmlElement('w:sectPr')
        body.append(sectPr)
    for old in sectPr.findall(qn('w:footerReference')):
        sectPr.remove(old)
    fr = OxmlElement('w:footerReference')
    fr.set(qn('w:type'), 'default')
    fr.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id', rId)
    sectPr.insert(0, fr)
    fr_first = OxmlElement('w:footerReference')
    fr_first.set(qn('w:type'), 'first')
    fr_first.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id', rId_first)
    sectPr.insert(0, fr_first)
    title_pg = sectPr.find(qn('w:titlePg'))
    if title_pg is None:
        title_pg = OxmlElement('w:titlePg')
        sectPr.append(title_pg)
    pg_num = sectPr.find(qn('w:pgNumType'))
    if pg_num is None:
        pg_num = OxmlElement('w:pgNumType')
        sectPr.append(pg_num)
    pg_num.set(qn('w:start'), '1')


def make_zadaniye_word(raw_text: str, d: dict) -> bytes:
    lv_ = 360
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21.001); sec.page_height   = Cm(29.700)
    sec.left_margin   = Cm(3.000);  sec.right_margin  = Cm(1.499)
    sec.top_margin    = Cm(2.000);  sec.bottom_margin = Cm(2.000)
    _init_heading(doc)
    _heading(doc, "ЗАДАНИЕ", lv_)
    p_tema = doc.add_paragraph()
    r_tema = p_tema.add_run(d["theme"])
    _sf(r_tema, 14, True)
    _spf(p_tema, "center", lv_, first_line=0, space_after=0)
    for ln in raw_text.splitlines():
        t = md_clean(ln)
        if t:
            _para(doc, t, size_pt=14, line=lv_, first_line=851, space_after=0)
    _add_page_numbers(doc)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()


def make_word(raw_text: str, d: dict) -> bytes:
    if d.get("service") == "zadaniye":
        return make_zadaniye_word(raw_text, d)
    spc    = spc_float(d.get("spacing","default"))
    secs   = int(d["sections"])
    parsed = parse_ai(raw_text, secs)
    lv_    = lv(spc)
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.001); sec.page_height = Cm(29.700)
    sec.left_margin = Cm(3.000); sec.right_margin = Cm(1.499)
    sec.top_margin  = Cm(2.000); sec.bottom_margin = Cm(2.000)
    _init_heading(doc)
    _copy_template_title(doc, d)
    _page_break(doc)
    _auto_toc(doc, lv_)
    _page_break(doc)
    _heading(doc, "Введение", lv_)
    for ln in parsed["intro"]:
        t = md_clean(ln)
        if t: _para(doc, t, size_pt=14, line=lv_, first_line=851, space_after=0)
    for ch in parsed["chapters"]:
        _page_break(doc)
        _heading(doc, ch["title"], lv_)
        for ln in ch["lines"]:
            t = md_clean(ln)
            if t: _para(doc, t, size_pt=14, line=lv_, first_line=851, space_after=0)
    _page_break(doc)
    _heading(doc, "ЗАКЛЮЧЕНИЕ", lv_)
    for ln in parsed["conclusion"]:
        t = md_clean(ln)
        if t: _para(doc, t, size_pt=14, line=lv_, first_line=851, space_after=0)
    _page_break(doc)
    _heading(doc, "Список литературы", lv_)
    for src in parsed["sources"]:
        t = md_clean(src)
        if not t or not re.match(r"^\d+\.", t): continue
        p = doc.add_paragraph(); r = p.add_run(t)
        _sf(r, 14, False); _spf(p, "both", lv_, first_line=851, space_after=0)
    _add_page_numbers(doc)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()


def t_progress(d: dict, pct: int, status: str) -> str:
    bar = "█"*(pct//10) + "░"*(10-pct//10)
    def _e(t): return str(t).replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
    return (f"⚙️ <b>{SVC_TM.get(d.get('service','referat'),'Iş')} taýarlanýar...</b>\n\n"
            f"📝 <i>{_e(d.get('theme',''))}</i> \n👤 {_e(d.get('fullname',''))}\n\n"
            f"<code>[{bar}]</code> <b>{pct}%</b>\n<i>{_e(status)}</i>")

def t_summary(d: dict) -> str:
    mug = d.get('teacher','—')
    tp  = d.get('teacher_position','')
    mug_line = f"{tp} {mug}".strip() if tp else mug
    def _e(t): return str(t).replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
    return (f"📋 <b>Sargyt jemi:</b>\n\n"
            f"📄 Görnüş : <b>{SVC_TM.get(d['service'],'?')}</b>\n"
            f"📖 Sapak  : {_e(d.get('subject','—'))}\n"
            f"📝 Tema   : <i>{_e(d.get('theme','—'))}</i>\n"
            f"🏫 Uni.   : {_e(d.get('university','—'))}\n"
            f"👤 Talyp  : {_e(d.get('fullname','—'))}\n"
            f"📚 Kurs   : {_e(d.get('course','—'))}-nji, {_e(d.get('group','—'))}\n"
            f"👨\u200d🏫 Mugallym: {_e(mug_line)}\n"
            f"📑 Bölüm  : {_e(d.get('sections','—'))}\n"
            f"📏 Aralyk : {_e(spc_str(d.get('spacing','default')))}\n"
            f"📄 Sahypa : {_e(d.get('pages','—'))}\n"
            f"🔗 Çeşme  : {_e(d.get('sources','—'))}\n")

async def send_file(uid: int, bot: Bot) -> None:
    if uid not in PAYMENT_PENDING:
        await bot.send_message(uid, "❌ Faýl tapylmady."); return
    info = PAYMENT_PENDING.pop(uid)
    d, fb = info["data"], info["bytes"]
    theme = re.sub(r"[^\w\s-]", "", d.get("theme", "work"))[:20].replace(" ", "_")
    fname = f"{d.get('service','work')}_{theme}.docx"
    await bot.send_document(uid, BufferedInputFile(fb, filename=fname),
        caption=(f"✅ <b>Işiňiz taýar!</b>\n\n"
                 f"📄 <code>{fname}</code>\n"
                 f"📝 {d.get('theme', '')}\n\n"
                 f"Üstünlik! 🎓   Täze sargyt: /start"),
        parse_mode="HTML")

async def deliver(uid: int, bot: Bot):
    if uid in CANCELLED_GENERATES:
        CANCELLED_GENERATES.discard(uid)
        PENDING.pop(uid, None)
        log.info(f"deliver: uid={uid} cancelled")
        return
    if uid not in PENDING:
        await bot.send_message(uid, "❌ Faýl tapylmady. Admin bilen habarlaşyň."); return
    info = PENDING.pop(uid)
    d    = info["data"]
    svc  = d.get("service", "referat")
    price = PRICE.get(svc, 300)
    svc_nm = {"referat":"Referat 📄","doklad":"Doklad 🎤","zadaniye":"Zadaniye 📝"}.get(svc, svc)
    PAYMENT_PENDING[uid] = info
    await bot.send_message(
        uid,
        f"✅ <b>{svc_nm} taýar boldy!</b>\n\n"
        f"📝 <i>{d.get('theme','')}</i>\n\n"
        f"💳 <b>Töleg maglumatlar:</b>\n"
        f"🏦 Kart    : <code>{CARD_NUMBER}</code>\n"
        f"📱 Telefon : <code>{PHONE_NUMBER}</code>\n"
        f"👤 At      : <b>{CARD_HOLDER}</b>\n"
        f"💰 Mukdar  : <b>{price} ₽</b>\n\n"
        f"📎 Töleg geçirensoň <b>çegiňizi ýa screenshotyňyzy</b> iberiň.\n"
        f"<i>Çek gelenden soň admin tassyklar we faýl size iberiler.</i>",
        parse_mode="HTML")


router = Router()


@router.callback_query(St.szad1, F.data == "req:yes")
async def hz_req_yes(cb: CallbackQuery, state: FSMContext):
    uid = cb.from_user.id
    REQ_ITEMS[uid] = []
    await state.update_data(has_req=True, req_items=[])
    await ask(cb,
        "✅ <b>Talaplы görnüş.</b>\n\n"
        "📎 Islendik zat iberiň — tekst, surat, faýl.\n"
        "Birnäçe habar iberip bilersiňiz.\n\n"
        "✅ Gutaransoň <b>«Talaplarym taýar»</b> düwmesine basyň:",
        KB_REQ_DONE)
    await state.set_state(St.szad2); await cb.answer()

@router.message(St.szad2)
async def hz_req_text(msg: Message, bot: Bot, state: FSMContext):
    uid = msg.from_user.id
    if uid not in REQ_ITEMS:
        REQ_ITEMS[uid] = []
    added = []
    txt = (msg.caption or msg.text or "").strip()
    if txt:
        REQ_ITEMS[uid].append({"type": "text", "content": txt}); added.append("tekst")
    if msg.photo:
        try:
            fobj = await bot.get_file(msg.photo[-1].file_id)
            fbuf = await bot.download_file(fobj.file_path)
            b64  = base64.b64encode(fbuf.read()).decode()
            REQ_ITEMS[uid].append({"type": "image", "mime": "image/jpeg", "content": b64})
            added.append("surat")
            await msg.answer(
                "🖼 <b>Surat kabul edildi!</b>\n\n"
                "⚠️ <i>Suradyň mazmunyny tekst bilen hem ýazyň — AI has gowy düşünýär.</i>\n\n"
                "<i>Mysal: «Şu suratdaky tablisany goş», «Şeýle usulda ýaz»</i>",
                parse_mode="HTML", reply_markup=KB_REQ_DONE)
            return
        except Exception as e:
            log.warning(f"Surat: {e}")
    if msg.document:
        try:
            fobj = await bot.get_file(msg.document.file_id)
            fbuf = await bot.download_file(fobj.file_path)
            raw  = fbuf.read(); mime = msg.document.mime_type or ""
            if "pdf" in mime:
                from pdfminer.high_level import extract_text as _ext
                import io as _io
                pt = _ext(_io.BytesIO(raw))
                REQ_ITEMS[uid].append({"type": "text", "content": f"[PDF: {msg.document.file_name}]\n{pt[:3000]}"})
            else:
                REQ_ITEMS[uid].append({"type": "text", "content": f"[{msg.document.file_name}]\n{raw.decode('utf-8','ignore')[:2000]}"})
            added.append("faýl")
        except Exception as e: log.warning(f"Faýl: {e}")
    count = len(REQ_ITEMS[uid])
    await msg.answer(
        f"✅ {'</b>, <b>'.join(added) if added else 'Kabul edildi'} ({count} zat)\n\n"
        "Ýene iberip bilersiňiz ýa-da:\n"
        "✅ <b>«Talaplarym taýar»</b> düwmesine basyň.",
        parse_mode="HTML", reply_markup=KB_REQ_DONE)

@router.callback_query(St.szad2, F.data == "req:done")
async def hz_req_done(cb: CallbackQuery, state: FSMContext):
    uid = cb.from_user.id
    items = REQ_ITEMS.pop(uid, [])
    texts = [i["content"] for i in items if i["type"] == "text"]
    req_text = "\n".join(texts) if texts else ""
    await state.update_data(req_text=req_text, req_items=items)
    await ask(cb,
        f"✅ Talaplar kabul edildi! ({len(items)} zat)\n\n"
        "📌 <b>2/2:</b> Zadaniýanyň temasyny ýazyň\n"
        "<i>Mysal: Maglumat howpsuzlygynyň esaslary</i>")
    await state.set_state(St.szad3); await cb.answer()

@router.callback_query(St.szad1, F.data == "req:no")
async def hz_req_no(cb: CallbackQuery, state: FSMContext):
    await state.update_data(has_req=False, req_text="")
    await ask(cb,
        "✅ Dowam edýär!\n\n"
        "📌 <b>2/2:</b> Zadaniýanyň temasyny ýazyň\n"
        "<i>Mysal: Maglumat howpsuzlygynyň esaslary</i>")
    await state.set_state(St.szad3); await cb.answer()

@router.message(St.szad3)
async def hz_tema(msg: Message, state: FSMContext):
    theme = msg.text.strip()
    if len(theme) < 5:
        await msg.answer("❌ Tema gysga! Iň az 5 harp ýazyň."); return
    uid_z = msg.from_user.id
    if uid_z in ACTIVE_GENERATES:
        await msg.answer("⏳ Sargydyňyz eýýäm taýarlanýar, sabyr ediň!", parse_mode="HTML"); return
    ACTIVE_GENERATES.add(uid_z)
    CANCELLED_GENERATES.discard(uid_z)
    ZADANIYE_GENERATING.add(uid_z)
    await state.update_data(
        theme=theme,
        subject="", university="", fullname="", course="1", group="",
        teacher="", teacher_position="",
        sections=2, spacing="default", pages=2, sources=0
    )
    d = await state.get_data()
    prog = await msg.answer(t_progress(d, 0, "Başlanýar..."), parse_mode="HTML")
    bot = msg.bot; cid = msg.chat.id; mid = prog.message_id
    async def pcb(pct, status):
        try: await bot.edit_message_text(t_progress(d, pct, status), chat_id=cid, message_id=mid, parse_mode="HTML")
        except: pass
    try:
        raw       = await call_deepseek(d, pcb)
        doc_bytes = await asyncio.get_running_loop().run_in_executor(None, make_word, raw, d)
        PENDING[uid_z] = {"bytes": doc_bytes, "data": d}
        await bot.edit_message_text(
            f"✅ <b>Zadaniýaňyz taýar boldy!</b>\n\n📝 <i>{theme}</i>",
            chat_id=cid, message_id=mid, parse_mode="HTML")
        await deliver(uid_z, bot)
    except Exception as exc:
        import traceback
        tb = traceback.format_exc()
        log.error(f"Zadaniye generate: {exc}\n{tb}")
        await bot.edit_message_text(
            f"❌ <b>Ýalňyşlyk!</b>\n\n<code>{str(exc)[:300]}</code>\n\nTäzeden başlamak: /start",
            chat_id=cid, message_id=mid, parse_mode="HTML")
    finally:
        ACTIVE_GENERATES.discard(uid_z)
        ZADANIYE_GENERATING.discard(uid_z)
        await state.clear()

@router.callback_query(F.data == "back:start")
async def h_back(cb: CallbackQuery, state: FSMContext):
    await state.clear()
    uid_b = cb.from_user.id
    PENDING.pop(uid_b, None)
    if uid_b not in ZADANIYE_GENERATING:
        ACTIVE_GENERATES.discard(uid_b)
        CANCELLED_GENERATES.add(uid_b)
    await cb.message.edit_text(
        "🎓 <b>Akademik Işler Boty</b>\n\n"
        "Salam! Referat, doklad we zadaniye taýarlaýaryn.\n"
        "Word faýly <b>rus dilinde</b> düzülýär.\n\n"
        "📌 <b>1/13:</b> Haýsy hyzmaty isleýärsiňiz?",
        parse_mode="HTML", reply_markup=KB_SVC)
    await state.set_state(St.s01)
    await cb.answer()

@router.message(CommandStart())
async def h_start(msg: Message, bot: Bot, state: FSMContext):
    await state.clear()
    uid = msg.from_user.id
    PENDING.pop(uid, None)
    if uid not in ZADANIYE_GENERATING:
        ACTIVE_GENERATES.discard(uid)
        CANCELLED_GENERATES.add(uid)
    if uid not in SEEN_USERS:
        SEEN_USERS.add(uid)
        if INTRO_VIDEO_URL:
            try:
                if INTRO_VIDEO_URL.startswith("http"):
                    await bot.send_message(uid,
                        f"🎓 <b>Akademik Işler Botyna hoş geldiňiz!</b>\n\n"
                        f"📹 Tanyşdyryş wideosyny görüň:\n{INTRO_VIDEO_URL}",
                        parse_mode="HTML")
                else:
                    await bot.send_video(uid, INTRO_VIDEO_URL,
                        caption="🎓 <b>Akademik Işler Botyna hoş geldiňiz!</b>",
                        parse_mode="HTML")
            except Exception as e:
                log.warning(f"Wideo iberilmedi: {e}")
    await msg.answer(
        "🎓 <b>Akademik Işler Boty</b>\n\n"
        "Salam! Referat, doklad we zadaniye taýarlaýaryn.\n"
        "Word faýly <b>rus dilinde</b> düzülýär.\n\n"
        "💰 <b>Bahalar:</b>\n"
        "📄 Referat  — <b>300 ₽</b>\n"
        "🎤 Doklad   — <b>300 ₽</b>\n"
        "📝 Zadaniye — <b>150 ₽</b>\n\n"
        "📌 <b>1/13:</b> Haýsy hyzmaty isleýärsiňiz?",
        parse_mode="HTML", reply_markup=KB_SVC)
    await state.set_state(St.s01)

@router.callback_query(St.s01, F.data.startswith("svc:"))
async def h01(cb: CallbackQuery, state: FSMContext):
    svc = cb.data.split(":")[1]; await state.update_data(service=svc)
    if svc == "zadaniye":
        await ask(cb,
            "✅ <b>Zadaniye 📝</b> saýlandy!\n\n"
            "📌 <b>1/2:</b> Ýörite talaplar barmy?\n\n"
            "• <b>Talaplы</b> — goşmaça talaplaryňyzy ýazyň\n"
            "• <b>Talapsyz</b> — ýok, dowam et",
            KB_REQ)
        await state.set_state(St.szad1)
    else:
        await ask(cb,
            f"✅ <b>{SVC_TM[svc]}</b> saýlandy!\n\n"
            "📌 <b>2/13:</b> Ýörite talaplar barmy?\n\n"
            "• <b>Talaplы</b> — öz talaplarыňyzy, faýl ýa-da surat ugradyp bilersiňiz\n"
            "• <b>Talapsyz</b> — adaty GOST görnüşi", KB_REQ)
        await state.set_state(St.s02)
    await cb.answer()

@router.callback_query(St.s02, F.data == "req:yes")
async def h02_yes(cb: CallbackQuery, state: FSMContext):
    uid = cb.from_user.id
    REQ_ITEMS[uid] = []
    await state.update_data(has_req=True, req_items=[])
    await ask(cb,
        "✅ <b>Talaplы görnüş.</b>\n\n"
        "📎 Islendik zat iberiň — tekst, surat, faýl, ýa hemmesini bilelikde.\n"
        "Birnäçe habar iberip bilersiňiz.\n\n"
        "✅ Gutaransoň <b>«Talaplarym taýar»</b> düwmesine basyň:",
        KB_REQ_DONE)
    await state.set_state(St.s02b); await cb.answer()

@router.message(St.s02b)
async def h02b(msg: Message, bot: Bot, state: FSMContext):
    uid = msg.from_user.id
    if uid not in REQ_ITEMS:
        REQ_ITEMS[uid] = []
    added = []
    txt = (msg.caption or msg.text or "").strip()
    if txt:
        REQ_ITEMS[uid].append({"type": "text", "content": txt})
        added.append("tekst")
    if msg.photo:
        try:
            fobj  = await bot.get_file(msg.photo[-1].file_id)
            fbuf  = await bot.download_file(fobj.file_path)
            raw   = fbuf.read()
            b64   = base64.b64encode(raw).decode()
            REQ_ITEMS[uid].append({"type": "image", "mime": "image/jpeg", "content": b64})
            added.append("surat")
            await msg.answer(
                "🖼 <b>Surat kabul edildi!</b>\n\n"
                "⚠️ <i>Suradyň mazmunyny ýa-da talaplaryny tekst bilen hem ýazyň — "
                "şeýdip AI has gowy düşünýär.</i>\n\n"
                "<i>Mysal: «Şu suratdaky tablisany goş», «Şeýle görnüşde ýaz» we ş.m.</i>",
                parse_mode="HTML", reply_markup=KB_REQ_DONE)
            return
        except Exception as e:
            log.warning(f"Surat okalmady: {e}")
    if msg.document:
        try:
            fobj  = await bot.get_file(msg.document.file_id)
            fbuf  = await bot.download_file(fobj.file_path)
            raw   = fbuf.read()
            mime  = msg.document.mime_type or "application/octet-stream"
            if "pdf" in mime:
                from pdfminer.high_level import extract_text as _ext
                import io as _io
                pdf_txt = _ext(_io.BytesIO(raw))
                if pdf_txt.strip():
                    REQ_ITEMS[uid].append({"type": "text",
                                           "content": f"[Faýldan: {msg.document.file_name}]\n{pdf_txt[:3000]}"})
                else:
                    REQ_ITEMS[uid].append({"type": "image", "mime": "image/jpeg",
                                           "content": base64.b64encode(raw).decode()})
            elif mime.startswith("text"):
                REQ_ITEMS[uid].append({"type": "text",
                                       "content": f"[Faýldan: {msg.document.file_name}]\n{raw.decode('utf-8','ignore')[:3000]}"})
            else:
                REQ_ITEMS[uid].append({"type": "text",
                                       "content": f"[{msg.document.file_name} faýly goşuldy]"})
            added.append("faýl")
        except Exception as e:
            log.warning(f"Faýl okalmady: {e}")
    count = len(REQ_ITEMS[uid])
    if added:
        await msg.answer(
            f"✅ <b>{'</b>, <b>'.join(added)}</b> kabul edildi! (Jemi: {count} zat)\n\n"
            "Ýene iberip bilersiňiz ýa-da:\n"
            "✅ <b>«Talaplarym taýar»</b> düwmesine basyň.",
            parse_mode="HTML", reply_markup=KB_REQ_DONE)
    else:
        await msg.answer("⚠️ Hiç zat tapylmady, täzeden iberiň.", reply_markup=KB_REQ_DONE)

@router.callback_query(St.s02b, F.data == "req:done")
async def h02b_done(cb: CallbackQuery, state: FSMContext):
    uid = cb.from_user.id
    items = REQ_ITEMS.pop(uid, [])
    texts = [i["content"] for i in items if i["type"] == "text"]
    req_text = "\n".join(texts) if texts else ""
    await state.update_data(req_text=req_text, req_items=items)
    await ask(cb,
        f"✅ Talaplar kabul edildi! ({len(items)} zat)\n\n"
        "📌 <b>3/13:</b> Uniwersitetiňiziň doly adyny ýazyň\n"
        "<i>Mysal: Türkmenistanyň Döwlet Energetika Instituty</i>")
    await state.set_state(St.s03); await cb.answer()

@router.callback_query(St.s02, F.data == "req:no")
async def h02_no(cb: CallbackQuery, state: FSMContext):
    await state.update_data(has_req=False, req_text="")
    await ask(cb,
        "✅ <b>Adaty GOST görnüşi.</b>\n\n"
        "📌 <b>3/13:</b> Uniwersitetiňiziň doly adyny ýazyň\n"
        "<i>Mysal: Türkmenistanyň Döwlet Energetika Instituty</i>")
    await state.set_state(St.s03); await cb.answer()

@router.message(St.s03)
async def h03(msg: Message, state: FSMContext):
    if len(msg.text.strip()) < 4: await msg.answer("❌ Iň az 4 harp ýazyň."); return
    await state.update_data(university=msg.text.strip())
    await msg.answer("✅ Kabul edildi!\n\n📌 <b>4/13:</b> Sapakyňyzyň adyny ýazyň\n"
                     "<i>Mysal: Ykdysadyýet nazaryýeti, Maglumat howpsuzlygy</i>\n\n"
                     "↩️ Täzeden başlamak: /start", parse_mode="HTML")
    await state.set_state(St.s04)

@router.message(St.s04)
async def h04(msg: Message, state: FSMContext):
    if len(msg.text.strip()) < 3: await msg.answer("❌ Iň az 3 harp ýazyň."); return
    await state.update_data(subject=msg.text.strip())
    await msg.answer(
        "✅ Kabul edildi!\n\n📌 <b>5/13:</b> Işiňiziň temasyny doly ýazyň <b>(HÖKMAN)</b>\n\n"
        "⚠️ <i>Tema anyk we doly bolmaly — AI şol temany şablona salar</i>\n\n"
        "<i>Mysal: Türkmenistanda ykdysady ösüşiň häzirki ýagdaýy</i>", parse_mode="HTML")
    await state.set_state(St.s05)

@router.message(St.s05)
async def h05(msg: Message, state: FSMContext):
    theme = msg.text.strip()
    if len(theme) < 8:
        await msg.answer("❌ Tema gysga! Iň az 8 harp.\n<i>Mysal: Türkmenistanda nebit-gaz senagatynyň ösüşi</i>",
                         parse_mode="HTML"); return
    await state.update_data(theme=theme)
    await msg.answer(f"✅ Tema: {theme}\n\n📌 <b>6/13:</b> Adyňyzy we Familiýaňyzy ýazyň\n"
                     "<i>Mysal: Myrat Mämmedow</i>", parse_mode="HTML")
    await state.set_state(St.s06)

@router.message(St.s06)
async def h06(msg: Message, state: FSMContext):
    if len(msg.text.strip()) < 3: await msg.answer("❌ Dogry ýazyň."); return
    await state.update_data(fullname=msg.text.strip())
    await msg.answer("✅ Kabul edildi!\n\n📌 <b>7/13:</b> Haýsy kursda okaýarsyňyz?",
                     parse_mode="HTML", reply_markup=KB_CRS)
    await state.set_state(St.s07)

@router.callback_query(St.s07, F.data.startswith("crs:"))
async def h07(cb: CallbackQuery, state: FSMContext):
    n = cb.data.split(":")[1]; await state.update_data(course=n)
    await ask(cb, f"✅ {n}-nji kurs!\n\n📌 <b>8/13:</b> Toparыňyzyň adyny ýazyň\n<i>Mysal: EHM-22, IT-21B</i>")
    await state.set_state(St.s08); await cb.answer()

@router.message(St.s08)
async def h08(msg: Message, state: FSMContext):
    await state.update_data(group=msg.text.strip())
    await msg.answer("✅ Kabul edildi!\n\n📌 <b>9/13:</b> Mugallymyň adyny ýazyň\n"
                     "<i>Mugallym ýok bolsa — «Geç» basyň</i>",
                     parse_mode="HTML", reply_markup=KB_SKIP)
    await state.set_state(St.s09)

@router.message(St.s09)
async def h09_text(msg: Message, state: FSMContext):
    teacher = msg.text.strip()
    await state.update_data(teacher=teacher)
    d = await state.get_data()
    if d.get("service") == "zadaniye":
        await state.update_data(teacher_position="", sections=2, spacing="default", pages=15, sources=10)
        d2 = await state.get_data()
        required_z = ["service","university","subject","theme","fullname","course","group"]
        if all(d2.get(k) for k in required_z):
            prog = await msg.answer(t_progress(d2, 0, "Başlanýar..."), parse_mode="HTML")
            bot = msg.bot; cid = msg.chat.id; mid = prog.message_id
            async def pcb(pct, status):
                try: await bot.edit_message_text(t_progress(d2, pct, status), chat_id=cid, message_id=mid, parse_mode="HTML")
                except: pass
            try:
                raw       = await call_deepseek(d2, pcb)
                doc_bytes = await asyncio.get_running_loop().run_in_executor(None, make_word, raw, d2)
                PENDING[msg.from_user.id] = {"bytes": doc_bytes, "data": d2}
                await bot.edit_message_text(f"✅ <b>Işiňiz taýar boldy!</b>\n\n{t_summary(d2)}",
                                             chat_id=cid, message_id=mid, parse_mode="HTML")
                await deliver(msg.from_user.id, bot)
            except Exception as exc:
                import traceback
                tb = traceback.format_exc()
                log.error(f"Generate zadaniye: {exc}\n{tb}")
                await bot.edit_message_text(
                    f"❌ <b>Ýalňyşlyk!</b>\n\n<code>{str(exc)[:300]}</code>\n\nTäzeden başlamak: /start",
                    chat_id=cid, message_id=mid, parse_mode="HTML")
            finally:
                await state.clear()
        return
    await msg.answer(
        "✅ Mugallym kabul edildi!\n\n📌 <b>9б/13:</b> Mugallymyň wezipesini ýazyň <b>(HÖKMAN)</b>\n"
        "<i>Mysal: доцент, профессор, ст. преподаватель</i>", parse_mode="HTML")
    await state.set_state(St.s09b)

@router.message(St.s09b)
async def h09b(msg: Message, state: FSMContext):
    if len(msg.text.strip()) < 2: await msg.answer("❌ Iň az 2 harp ýazyň."); return
    await state.update_data(teacher_position=msg.text.strip())
    await msg.answer("✅ Kabul edildi!\n\n📌 <b>10/13:</b> Näçe esasy bölüm bolmaly?",
                     parse_mode="HTML", reply_markup=KB_SEC)
    await state.set_state(St.s10)

@router.callback_query(St.s09, F.data == "skip:teacher")
async def h09_skip(cb: CallbackQuery, state: FSMContext):
    await state.update_data(teacher="", teacher_position="")
    d = await state.get_data()
    if d.get("service") == "zadaniye":
        await state.update_data(sections=2, spacing="default", pages=15, sources=10)
        d2 = await state.get_data()
        await _run_generate(cb, state, d2)
        return
    await ask(cb, "✅ Geçildi!\n\n📌 <b>10/13:</b> Näçe esasy bölüm bolmaly?", KB_SEC)
    await state.set_state(St.s10); await cb.answer()

@router.callback_query(St.s10, F.data.startswith("sec:"))
async def h10(cb: CallbackQuery, state: FSMContext):
    n = int(cb.data.split(":")[1]); await state.update_data(sections=n)
    d = await state.get_data()
    if d.get("service") == "zadaniye":
        await _run_generate(cb, state, d)
        return
    await ask(cb,
        f"✅ {n} bölüm!\n\n📌 <b>11/13:</b> Setirler aralygy näçe?\n\n"
        "• <b>Adaty (1.5)</b> — GOST standart\n• <b>Üýtgetmek</b> — 1.0 / 1.25 / 1.5", KB_SPC)
    await state.set_state(St.s11); await cb.answer()

@router.callback_query(St.s11, F.data == "spc:default")
async def h11_def(cb: CallbackQuery, state: FSMContext):
    await state.update_data(spacing="default")
    await ask(cb,
        "✅ Setirler aralygy <b>1.5</b> saýlandy!\n\n📌 <b>12/13:</b> Näçe sahypa?\n\n"
        "💡 <i>10—20 maslahat</i>\nSan ýazyň <i>(mysal: 15)</i>:")
    await state.set_state(St.s12); await cb.answer()

@router.callback_query(St.s11, F.data == "spc:custom")
async def h11_cus(cb: CallbackQuery, state: FSMContext):
    await ask(cb,
        "✏️ <b>Setirler aralygyny saýlaň:</b>\n\n"
        "• <b>1.0</b> — gysga\n• <b>1.25</b> — orta\n• <b>1.5</b> — GOST", KB_SPCV)
    await state.set_state(St.s11b); await cb.answer()

@router.callback_query(St.s11b, F.data.startswith("spv:"))
async def h11b(cb: CallbackQuery, state: FSMContext):
    val = cb.data.split(":")[1]; await state.update_data(spacing=val)
    await ask(cb,
        f"✅ Setirler aralygy <b>{val}</b> saýlandy!\n\n📌 <b>12/13:</b> Näçe sahypa?\n\n"
        "💡 <i>10—20 maslahat</i>\nSan ýazyň <i>(mysal: 15)</i>:")
    await state.set_state(St.s12); await cb.answer()

@router.message(St.s12)
async def h12(msg: Message, state: FSMContext):
    try: n = int(msg.text.strip())
    except: await msg.answer("❌ Diňe san ýazyň! <i>Mysal: 15</i>", parse_mode="HTML"); return
    if not (10 <= n <= 20): await msg.answer("⚠️ Diňe 10—20 arasynda!", parse_mode="HTML"); return
    await state.update_data(pages=n)
    await msg.answer(
        "✅ Kabul edildi!\n\n📌 <b>13/13:</b> «Список литературы» näçe çeşme?\n"
        "<i>8—20 arasynda saýlaň</i>", parse_mode="HTML", reply_markup=kb_src())
    await state.set_state(St.s13)

async def _run_generate(cb: CallbackQuery, state: FSMContext, d: dict = None):
    if d is None:
        d = await state.get_data()
    if d.get("service") == "zadaniye":
        required = ["service","university","subject","theme","fullname","course","group"]
    else:
        required = ["service","university","subject","theme","fullname","course","group","sections","pages","sources"]
    missing = [k for k in required if not d.get(k)]
    if missing:
        await cb.answer("❌ Ýetmeýän maglumat!", show_alert=True)
        await cb.message.answer(
            f"❌ Ýetmeýänler: <code>{chr(44).join(missing)}</code>\nTäzeden başlamak: /start",
            parse_mode="HTML")
        await state.clear(); return
    uid_gen = cb.from_user.id
    if uid_gen in ACTIVE_GENERATES:
        await cb.answer("⏳ Sargydyňyz eýýäm taýarlanýar, sabyr ediň!", show_alert=True); return
    ACTIVE_GENERATES.add(uid_gen)
    CANCELLED_GENERATES.discard(uid_gen)
    prog = await cb.message.edit_text(t_progress(d, 0, "Başlanýar..."), parse_mode="HTML")
    await state.set_state(St.s13); await cb.answer()
    bot = cb.bot; cid = cb.message.chat.id; mid = prog.message_id
    async def pcb(pct, status):
        try: await bot.edit_message_text(t_progress(d, pct, status), chat_id=cid, message_id=mid, parse_mode="HTML")
        except: pass
    try:
        raw       = await call_deepseek(d, pcb)
        doc_bytes = await asyncio.get_running_loop().run_in_executor(None, make_word, raw, d)
        PENDING[cb.from_user.id] = {"bytes": doc_bytes, "data": d}
        await bot.edit_message_text(f"✅ <b>Işiňiz taýar boldy!</b>\n\n{t_summary(d)}",
                                     chat_id=cid, message_id=mid, parse_mode="HTML")
        await deliver(cb.from_user.id, bot)
    except Exception as exc:
        import traceback
        tb = traceback.format_exc()
        log.error(f"Generate: {exc}\n{tb}")
        await bot.edit_message_text(
            f"❌ <b>Ýalňyşlyk!</b>\n\n<code>{str(exc)[:300]}</code>\n\nTäzeden başlamak: /start",
            chat_id=cid, message_id=mid, parse_mode="HTML")
    finally:
        ACTIVE_GENERATES.discard(uid_gen)
        await state.clear()

@router.callback_query(St.s13, F.data.startswith("src:"))
async def h13_generate(cb: CallbackQuery, state: FSMContext):
    n = int(cb.data.split(":")[1]); await state.update_data(sources=n)
    d = await state.get_data()
    await _run_generate(cb, state, d)


# ✅ DÜZEDIŞ #3: F.video handler — admin üçin hem, ulanyjy üçin hem ikisini birleşdirdim
@router.message(F.video | F.video_note)
async def h_video(msg: Message):
    uid = msg.from_user.id
    # Admin bolsa — file_id ber
    if uid in ADMIN_IDS:
        fid = msg.video.file_id if msg.video else msg.video_note.file_id
        await msg.answer(
            f"✅ <b>Wideo file_id:</b>\n\n<code>{fid}</code>\n\n"
            f"Şony <code>INTRO_VIDEO_URL</code>-e ýaz.",
            parse_mode="HTML")
        return
    # Adaty ulanyjy we töleg garaşýar
    if uid not in PAYMENT_PENDING:
        return
    await msg.answer(
        "❌ <b>Wideo kabul edilmeýär!</b>\n\n"
        "Çek hökmünde diňe şulary iberiň:\n"
        "📸 Surat (screenshot)\n"
        "📄 Faýl (PDF, PNG, JPG)\n\n"
        "Töleg screenshotyny ýa-da çek suratyny iberiň:",
        parse_mode="HTML")

@router.message(F.photo | F.document)
async def h_receipt(msg: Message, bot: Bot, state: FSMContext):
    uid = msg.from_user.id
    cur = await state.get_state()
    if cur in (St.s02b.state, St.szad2.state):
        return
    if uid not in PAYMENT_PENDING:
        return
    log.info(f"h_receipt: uid={uid} çek iberdi")
    d     = PAYMENT_PENDING[uid]["data"]
    svc   = d.get("service", "referat")
    price = PRICE.get(svc, 300)
    svc_nm = {"referat":"Referat 📄","doklad":"Doklad 🎤","zadaniye":"Zadaniye 📝"}.get(svc, svc)
    confirm_kb = InlineKeyboardMarkup(inline_keyboard=[[
        InlineKeyboardButton(text="✅ Tassykla", callback_data=f"confirm:{uid}"),
        InlineKeyboardButton(text="❌ Ret et",   callback_data=f"reject:{uid}"),
    ]])
    caption = (f"🧾 <b>Çek geldi — tassykla</b>\n\n"
               f"👤 <code>{uid}</code> | {svc_nm} <b>{price} ₽</b>\n"
               f"📝 <i>{d.get('theme','—')[:50]}</i>\n"
               f"👥 {d.get('fullname','—')} | {d.get('course','—')}-kurs")
    sent = False
    for admin in ADMIN_IDS:
        try:
            if msg.photo:
                await bot.send_photo(admin, msg.photo[-1].file_id,
                                     caption=caption, parse_mode="HTML",
                                     reply_markup=confirm_kb)
                sent = True
            elif msg.document:
                await bot.send_document(admin, msg.document.file_id,
                                        caption=caption, parse_mode="HTML",
                                        reply_markup=confirm_kb)
                sent = True
        except Exception as e:
            log.error(f"Admine {admin} çek ibermek: {e}")
    if sent:
        await msg.answer("✅ <b>Çegiňiz alyndy!</b>\n\nAdmin barlaýar, az garaşyň...",
                         parse_mode="HTML")
    else:
        await msg.answer("⚠️ Çek iberilmedi, täzeden synanyşyň.", parse_mode="HTML")

@router.callback_query(F.data.startswith("confirm:"))
async def h_confirm(cb: CallbackQuery, bot: Bot):
    if cb.from_user.id not in ADMIN_IDS:
        await cb.answer("❌ Diňe admin!", show_alert=True); return
    uid = int(cb.data.split(":")[1])
    if uid not in PAYMENT_PENDING:
        await cb.answer("⚠️ Faýl tapylmady ýa-da eýýäm iberildi", show_alert=True)
        await cb.message.edit_reply_markup(reply_markup=None); return
    old_text = cb.message.text or cb.message.caption or ""
    new_text = old_text + "\n\n✅ <b>TASSYKLANDI</b>"
    try:
        if cb.message.text:
            await cb.message.edit_text(new_text, parse_mode="HTML", reply_markup=None)
        else:
            await cb.message.edit_caption(new_text, parse_mode="HTML", reply_markup=None)
    except Exception: pass
    await cb.answer("✅ Tassyklandy, faýl iberilýär")
    await send_file(uid, bot)
    try:
        await bot.send_message(uid, "✅ <b>Tölegiňiz tassyklandy!</b> Faýlyňyz häzir iberilýär...",
                               parse_mode="HTML")
    except Exception: pass

@router.callback_query(F.data.startswith("reject:"))
async def h_reject(cb: CallbackQuery, bot: Bot):
    if cb.from_user.id not in ADMIN_IDS:
        await cb.answer("❌ Diňe admin!", show_alert=True); return
    uid = int(cb.data.split(":")[1])
    info_r = PAYMENT_PENDING.pop(uid, None)
    old_text = cb.message.text or cb.message.caption or ""
    new_text = old_text + "\n\n❌ <b>RET EDILDI</b>"
    try:
        if cb.message.text:
            await cb.message.edit_text(new_text, parse_mode="HTML", reply_markup=None)
        else:
            await cb.message.edit_caption(new_text, parse_mode="HTML", reply_markup=None)
    except Exception: pass
    await cb.answer("❌ Ret edildi")
    try:
        svc   = info_r["data"].get("service","referat") if info_r else "referat"
        price = PRICE.get(svc, 300)
        await bot.send_message(
            uid,
            f"❌ <b>Tölegiňiz tassyklanmady.</b>\n\n"
            f"Dogry mukdarda ({price} ₽) töleg geçirip,\n"
            f"çegiňizi täzeden iberiň ýa-da admin bilen habarlaşyň.",
            parse_mode="HTML")
    except Exception: pass

@router.message(F.text.regexp(r"^/send\s+\d+"))
async def a_send(msg: Message, bot: Bot):
    if msg.from_user.id not in ADMIN_IDS: return
    uid = int(msg.text.split()[1]); await deliver(uid, bot)
    await msg.answer(f"✅ {uid} iberildi.")

@router.message(F.text == "/orders")
async def a_orders(msg: Message):
    if msg.from_user.id not in ADMIN_IDS: return
    if not PENDING: await msg.answer("📭 Garaşyp duran sargyt ýok."); return
    lines = ["📋 <b>Sargytlar:</b>\n"]
    for uid, info in PENDING.items():
        d = info["data"]
        lines.append(f"👤 <code>{uid}</code> | <i>{d.get('theme','?')[:25]}</i> | {d.get('pages','?')} sah.\n   /send {uid}\n")
    await msg.answer("\n".join(lines), parse_mode="HTML")

@router.message(F.text == "/admin")
async def a_help(msg: Message):
    if msg.from_user.id not in ADMIN_IDS: return
    await msg.answer(
        "🔧 <b>Admin:</b>\n"
        "<code>/orders</code> — sargytlar\n"
        "<code>/send &lt;id&gt;</code> — iber\n"
        "<code>/getfileid</code> — wideo iberip file_id al",
        parse_mode="HTML")

@router.message(F.text == "/getfileid")
async def a_getfileid(msg: Message):
    if msg.from_user.id not in ADMIN_IDS: return
    await msg.answer("📹 Indi wideoňy şu söhbete iber — file_id-ni bererin.")


async def main():
    from aiogram.client.default import DefaultBotProperties
    from aiogram.client.session.aiohttp import AiohttpSession

    session = AiohttpSession(timeout=300)
    bot = Bot(token=BOT_TOKEN, session=session,
              default=DefaultBotProperties(parse_mode="HTML"))
    me  = await bot.get_me()
    log.info(f"✅ @{me.username} işe başlady!")

    # ✅ DÜZEDIŞ #5: Redis URL env-den, hardcode aýryldy
    storage = None
    if REDIS_URL and _HAS_REDIS:
        try:
            storage = RedisStorage.from_url(REDIS_URL)
            log.info("✅ Redis storage ulanylýar")
        except Exception as e:
            log.warning(f"⚠️ Redis ulanylmady ({e}), MemoryStorage-a geçýär")
            storage = MemoryStorage()
    else:
        storage = MemoryStorage()
        log.info("⚠️  MemoryStorage ulanylýar (bot ýapylsa state ýitýär)")

    dp  = Dispatcher(storage=storage)
    dp.include_router(router)

    while True:
        try:
            await dp.start_polling(
                bot,
                allowed_updates=["message","callback_query"],
                polling_timeout=30,
            )
        except Exception as e:
            log.error(f"Polling error: {e} — 5 sek garaşyp täzeden başlanýar")
            await asyncio.sleep(5)


if __name__ == "__main__":
    asyncio.run(main())
